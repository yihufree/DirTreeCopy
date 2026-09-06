import os
import sys
import re  # 20260906 223308 新增：解析用户配置中保存的窗口几何字符串
import json  # 20260906 201500 新增：用于保存/读取用户界面偏好（目录树列宽）
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, font
import shutil
from operation_history import OperationHistory

APP_NAME = "目录和文件复制及重命名工具"
APP_VERSION = "V1.5"  # 20260906 192644 版本升级至 V1.5（新增关于按钮、多维重命名勾选式、全层级目录名称修改）
APP_BUILD_DATE = "20260906"  # 20260906 192644 构建日期更新为 20260906
APP_RELEASE_DATE = "2026/09/06"  # 20260906 192644 更新日期更新为 2026/09/06
APP_DEVELOP_DATE = "2026/09/06"  # 20260906 192644 开发日期更新为 2026/09/06
APP_AUTHOR = "飞歌"
APP_EXECUTABLE_NAME = f"DirCopyTool_{APP_BUILD_DATE}_{APP_VERSION}"
APP_WINDOW_TITLE = f"{APP_NAME}   {APP_BUILD_DATE}  {APP_AUTHOR}"

# 20260906 201500 新增：目录树列宽拖动与用户偏好配置相关常量
CONFIG_FILE_NAME = ".dirtreecopy_config.json"  # 用户偏好配置文件（保存在用户主目录）
TREE_COLUMN_NAMES = ("#0", "checked", "type", "size")  # 目录树列标识
TREE_COLUMN_MIN_WIDTH = 40  # 目录树列最小宽度（像素）
TREE_COLUMN_MAX_WIDTH = 800  # 目录树列最大宽度（像素）

class DirCopyApp:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_WINDOW_TITLE)
        
        # 获取屏幕尺寸
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        
        # 动态计算窗口大小
        base_window_width = max(800, min(1200, int(screen_width * 0.60)))  # 20260402 120300 以屏幕宽度的60%作为基础宽度
        window_width = max(720, base_window_width - 150)  # 20260402 120000 默认宽度再减小150（不小于最小宽度）
        window_height = 650
        
        # 计算窗口位置（水平居中，垂直距离上边50像素）
        x = (screen_width - window_width) // 2
        y = 50
        
        # 确保窗口下边框不低于任务栏上边框
        taskbar_height = 40  # 估计任务栏高度
        max_y = screen_height - window_height - taskbar_height
        if y > max_y:
            y = max_y
        
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # 设置窗口最小尺寸
        self.root.minsize(720, 500)  # 20260402 112200 允许窗口更紧凑但保留可用空间
        
        # 设置默认字体大小（缩小为原来的80%）
        default_font = font.nametofont("TkDefaultFont")
        default_font.configure(size=10)  # 20260402 101800 界面字体整体降低一级（不含标题与列表框）
        
        text_font = font.nametofont("TkTextFont")
        text_font.configure(size=10)  # 20260402 101800 界面字体整体降低一级（不含标题与列表框）
        
        fixed_font = font.nametofont("TkFixedFont")
        fixed_font.configure(size=10)  # 20260402 101800 界面字体整体降低一级（不含标题与列表框）
        
        # 创建按钮专用的较大字体
        self.button_font = font.Font(family="TkDefaultFont", size=11, weight="bold")  # 20260402 101800 界面字体整体降低一级（不含标题与列表框）
        
        # 创建软件标题字体（二号字体约18磅）
        self.title_font = font.Font(family="华文新魏", size=20, weight="bold")
        
        # 创建加粗标签字体
        self.bold_label_font = font.Font(family="TkDefaultFont", size=10, weight="bold")  # 20260402 101800 界面字体整体降低一级（不含标题与列表框）
        
        # 初始化变量
        self.source_dir = tk.StringVar()
        self.dest_dir = tk.StringVar()
        self.copy_mode = tk.StringVar(value="custom")
        self.tree_display_mode = tk.StringVar(value="normal")
        self.export_format = tk.StringVar(value="txt")
        self.show_directory_sizes = tk.BooleanVar(value=False)  # 20260402 085901 增加目录大小显示开关（默认关闭以提升大目录性能）
        self.directory_size_cache = {}
        self.runtime_warnings = []
        self.ui_update_interval = 50
        self.ui_update_counter = 0
        self.preview_filter_mode = "all"
        self.tree_item_paths = {}  # 20260402 085901 目录树节点路径映射，用于懒加载与可靠路径解析
        self.lazy_loaded_items = set()  # 20260402 085901 记录已完成懒加载的节点，避免重复加载
        self.lazy_placeholder_text = "..."  # 20260402 085901 懒加载占位节点文本
        self.current_tree_display_mode_override = None  # 20260402 085901 记录当前刷新时的显示模式覆盖值
        # 20260906 201500 新增：目录树表头列宽拖动状态
        self._tree_resizing = False
        self._tree_resize_col = None
        self._tree_resize_start_x = 0
        self._tree_resize_start_w = 0
        
        # 存储复选框状态
        self.checked_items = set()
        
        # 初始化操作历史管理器
        self.operation_history = OperationHistory(max_history_size=50)
        self._setup_backup_directory()
        
        # 创建主界面
        self.create_main_interface()
        # 20260906 223308 新增：应用用户记忆的主窗口位置/大小，并绑定关闭时保存（窗口位置/大小持久化）
        self._apply_saved_window_geometry()
        self.root.protocol("WM_DELETE_WINDOW", self._on_main_window_close)
        self.task_runner = None  # 20260402 091600 初始化任务调度器引用
        self.task_running = False  # 20260402 091600 任务运行状态标志
        self.cancel_requested = False  # 20260402 091600 取消标志
        self._test_force_sync = False  # 20260402 091600 测试专用：强制同步执行任务，便于无主循环用例
        
    def create_rename_dialog(self, title):
        """创建重命名对话框的通用方法"""
        dialog = tk.Toplevel(self.root)
        dialog.title(title)
        dialog.geometry("400x200")
        
        dialog_font = ('TkDefaultFont', 9)  # 20260402 101800 界面字体整体降低一级（不含标题与列表框）
        
        ttk.Label(dialog, text="要查找的字符串:", font=dialog_font).pack(pady=5)
        find_var = tk.StringVar()
        ttk.Entry(dialog, textvariable=find_var, font=dialog_font).pack(pady=5)
        
        ttk.Label(dialog, text="替换为:", font=dialog_font).pack(pady=5)
        replace_var = tk.StringVar()
        ttk.Entry(dialog, textvariable=replace_var, font=dialog_font).pack(pady=5)
        
        return dialog, find_var, replace_var
    
    def _setup_backup_directory(self):
        """设置备份目录"""
        app_data_dir = os.getenv("LOCALAPPDATA")  # 20260402 085901 将备份目录迁移到用户可写目录，兼容PyInstaller单文件运行环境
        if not app_data_dir:
            app_data_dir = os.path.expanduser(r"~\AppData\Local")  # 20260402 085901 LOCALAPPDATA缺失时使用兜底路径
        backup_dir = os.path.join(app_data_dir, "DirTreeCopy", "backup")  # 20260402 085901 统一备份目录位置
        try:
            os.makedirs(backup_dir, exist_ok=True)  # 20260402 085901 确保备份目录存在
            self.operation_history.set_backup_directory(backup_dir)  # 20260402 085901 使用用户目录作为备份位置
        except Exception:
            current_dir = os.path.dirname(os.path.abspath(__file__))  # 20260402 085901 目录创建失败时回退到程序目录
            fallback_dir = os.path.join(current_dir, "backup")  # 20260402 085901 回退备份目录
            self.operation_history.set_backup_directory(fallback_dir)  # 20260402 085901 回退策略

    def _process_pending_ui(self, force=False):
        if force:
            self.ui_update_counter = 0
            self.root.update_idletasks()
            return

        self.ui_update_counter += 1
        if self.ui_update_counter >= self.ui_update_interval:
            self.ui_update_counter = 0
            self.root.update_idletasks()

    def _reset_runtime_warnings(self):
        self.runtime_warnings = []

    def _add_runtime_warning(self, message):
        if message and message not in self.runtime_warnings:
            self.runtime_warnings.append(message)

    def _show_runtime_warnings(self, title="提示"):
        if not self.runtime_warnings:
            return

        max_visible = 8
        visible_messages = self.runtime_warnings[:max_visible]
        remaining_count = len(self.runtime_warnings) - len(visible_messages)
        warning_text = "\n".join(f"• {message}" for message in visible_messages)
        if remaining_count > 0:
            warning_text += f"\n• 另有 {remaining_count} 条类似提示未展开显示"
        messagebox.showwarning(title, warning_text)

    def _show_info_message(self, title, message):
        messagebox.showinfo(title, message)

    def _show_warning_message(self, title, message):
        messagebox.showwarning(title, message)

    def _show_error_message(self, title, message):
        messagebox.showerror(title, message)

    def _require_source_directory(self):
        if self.source_dir.get():
            return True

        self._show_warning_message("警告", "请先选择源目录!")
        return False

    def _warn_check_mode_required(self):  # 20260906 201500 新增：非勾选模式统一提示
        """当前操作模式不支持勾选时，提示用户切换到支持勾选的操作模式"""
        self._show_warning_message(
            "提示",
            "当前操作模式不支持勾选选择。\n\n"
            "请先切换到“复制选定目录和文件”或“复制选定层级目录”模式，\n"
            "然后浏览源目录并勾选需要处理的目录/文件。"
        )

    def _finalize_tree_change(self, success_title, success_message, refresh_mode=None, warning_title=None, update_history=False):
        self._show_info_message(success_title, success_message)
        if warning_title:
            self._show_runtime_warnings(warning_title)
        if refresh_mode is not None:
            self.refresh_tree(refresh_mode)
        if update_history:
            self.update_history_buttons()

    def _execute_history_action(self, action_name, success_message, failure_message, dialog=None):
        action = getattr(self.operation_history, action_name)
        if action():
            self._show_info_message("成功", success_message)
            self.refresh_tree()
            self.update_history_buttons()
            if dialog is not None:
                dialog.destroy()
                self.show_operation_history()
            return True

        self._show_warning_message("警告", failure_message)
        return False

    def _get_source_items(self, target_type):
        source_path = self.source_dir.get()
        predicate = os.path.isdir if target_type == "directory" else os.path.isfile

        try:
            items = []
            with os.scandir(source_path) as iterator:  # 20260402 084500 使用scandir减少目录遍历系统调用次数
                for entry in iterator:
                    if target_type == "directory":
                        if entry.is_dir(follow_symlinks=False):
                            items.append(entry.name)
                    else:
                        if entry.is_file(follow_symlinks=False):
                            items.append(entry.name)
            return items
        except Exception as e:
            self._show_error_message("错误", f"无法读取目录: {str(e)}")
            return None

    def _get_control_value(self, control):
        return control.get() if hasattr(control, "get") else control

    def _resolve_multi_rename_controls(self, controls, target_type):
        if controls is not None:
            return controls

        prefix = "file_" if target_type == "file" else ""
        return {
            "prefix_num_var": getattr(self, f"{prefix}prefix_num_var"),
            "prefix_conn_var": getattr(self, f"{prefix}prefix_conn_var"),
            "prefix_text_var": getattr(self, f"{prefix}prefix_text_var"),
            "suffix_text_var": getattr(self, f"{prefix}suffix_text_var"),
            "suffix_conn_var": getattr(self, f"{prefix}suffix_conn_var"),
            "suffix_num_var": getattr(self, f"{prefix}suffix_num_var"),
            "prefix_num_combo": getattr(self, f"{prefix}prefix_num_combo"),
            "prefix_conn_combo": getattr(self, f"{prefix}prefix_conn_combo"),
            "prefix_text_entry": getattr(self, f"{prefix}prefix_text_entry"),
            "suffix_text_entry": getattr(self, f"{prefix}suffix_text_entry"),
            "suffix_conn_combo": getattr(self, f"{prefix}suffix_conn_combo"),
            "suffix_num_combo": getattr(self, f"{prefix}suffix_num_combo"),
        }

    def _build_multi_rename_name(self, base_name, index, controls, extension=""):
        prefix_parts = []
        suffix_parts = []

        if self._get_control_value(controls["prefix_num_var"]):
            num_type = self._get_control_value(controls["prefix_num_combo"])
            prefix_parts.append(self.generate_sequence_number(index, num_type))

        if self._get_control_value(controls["prefix_conn_var"]):
            prefix_parts.append(self._get_control_value(controls["prefix_conn_combo"]))

        if self._get_control_value(controls["prefix_text_var"]):
            prefix_text = self._get_control_value(controls["prefix_text_entry"]).strip()
            if prefix_text:
                prefix_parts.append(prefix_text)

        if self._get_control_value(controls["suffix_text_var"]):
            suffix_text = self._get_control_value(controls["suffix_text_entry"]).strip()
            if suffix_text:
                suffix_parts.append(suffix_text)

        if self._get_control_value(controls["suffix_conn_var"]):
            suffix_parts.append(self._get_control_value(controls["suffix_conn_combo"]))

        if self._get_control_value(controls["suffix_num_var"]):
            num_type = self._get_control_value(controls["suffix_num_combo"])
            suffix_parts.append(self.generate_sequence_number(index, num_type))

        return "".join(prefix_parts) + base_name + "".join(suffix_parts) + extension

    def _prepare_rename_target(self, parent_path, old_name, new_name, item_label):
        if not new_name or not new_name.strip():
            return None, f"{item_label} '{old_name}' 的新名称为空，已跳过"

        if new_name == old_name:
            return None, None

        old_path = os.path.join(parent_path, old_name)
        new_path = os.path.join(parent_path, new_name)

        if os.path.normcase(old_path) != os.path.normcase(new_path) and os.path.exists(new_path):
            return None, f"{item_label} '{new_name}' 已存在，跳过重命名 '{old_name}'"

        return new_path, None

    def _build_rename_summary_message(self, type_name, matched_count, success_count, skipped_count, failed_count):
        return f"重命名完成！匹配 {matched_count} 个{type_name}，成功 {success_count} 个，跳过 {skipped_count} 个，失败 {failed_count} 个。"

    def _format_size_in_kb(self, size_bytes):
        size_kb = size_bytes / 1024
        return f"{size_kb:.1f}" if size_kb >= 0.1 else "0.1"

    def _build_directory_size_cache(self, root_path, display_mode):
        self.directory_size_cache = {}
        if not root_path or display_mode == "files_only":
            return

        # 20260828 111834 改用 os.scandir + entry.stat 计算目录大小，减少路径拼接与系统调用
        try:
            self._scan_directory_sizes(root_path)
        except (OSError, IOError, RecursionError):
            self.directory_size_cache = {}

    def _scan_directory_sizes(self, root_path):
        """迭代式计算目录大小并写入缓存（自底向上，使用 scandir 的 entry.stat）

        20260906 223308 由递归改写为显式栈迭代式，彻底避免深层目录触发 RecursionError。
        分两阶段：先沿树收集全部目录（每目录仅 scandir 一次），再逆序累加子目录大小（先子后父）。
        """
        # 阶段一：显式栈深度优先收集目录，并记录每个目录的直接文件大小与子目录列表
        dir_file_sizes = {}  # 目录自身的直接文件大小（不含子目录）
        dir_children = {}    # 目录的直接子目录列表
        dirs = []            # 已收集目录（父先子后）
        stack = [root_path]
        while stack:
            path = stack.pop()
            try:
                with os.scandir(path) as it:
                    entries = list(it)
            except (PermissionError, OSError):
                # 无法访问的目录按 0 计入缓存（与原递归行为一致）
                dirs.append(path)
                dir_file_sizes[path] = 0
                dir_children[path] = []
                continue
            dirs.append(path)
            file_total = 0
            children = []
            for entry in entries:
                try:
                    if entry.is_dir(follow_symlinks=False):
                        children.append(entry.path)
                        stack.append(entry.path)
                    else:
                        file_total += entry.stat(follow_symlinks=False).st_size
                except (OSError, IOError):
                    continue
            dir_file_sizes[path] = file_total
            dir_children[path] = children
            self._process_pending_ui()
        # 阶段二：逆序（先深层后浅层）把子目录大小累加到父目录并写入缓存
        for path in reversed(dirs):
            total_size = dir_file_sizes.get(path, 0)
            for sub_path in dir_children.get(path, ()):
                total_size += self.directory_size_cache.get(sub_path, 0)
            self.directory_size_cache[path] = total_size
    
    def undo_last_operation(self):
         """撤销上一个操作"""
         self._execute_history_action("undo", "操作已撤销", "无法撤销当前操作")
    
    def redo_last_operation(self):
         """重做上一个操作"""
         self._execute_history_action("redo", "操作已重做", "无法重做操作")
    
    def show_operation_history(self):
        """显示操作历史对话框"""
        history_dialog = tk.Toplevel(self.root)
        history_dialog.title("操作历史")
        history_dialog.geometry("600x400")
        
        # 设置对话框位置
        history_dialog.update_idletasks()
        screen_width = history_dialog.winfo_screenwidth()
        dialog_width = 600
        x = (screen_width - dialog_width) // 2
        y = 80
        history_dialog.geometry(f"{dialog_width}x400+{x}+{y}")
        
        history_dialog.transient(self.root)
        history_dialog.grab_set()
        
        main_frame = ttk.Frame(history_dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(main_frame, text="操作历史记录", 
                               font=("TkDefaultFont", 12, "bold"))
        title_label.pack(pady=(0, 10))
        
        # 按钮框架 - 先创建按钮框架确保显示
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0), side=tk.BOTTOM)

        # 创建树形视图显示历史记录
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        history_tree = ttk.Treeview(tree_frame, columns=("time", "operation", "status"), show="headings")
        history_tree.heading("time", text="时间")
        history_tree.heading("operation", text="操作")
        history_tree.heading("status", text="状态")

        history_tree.column("time", width=150)
        history_tree.column("operation", width=300)
        history_tree.column("status", width=100)

        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=history_tree.yview)
        history_tree.configure(yscrollcommand=v_scrollbar.set)

        history_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 填充历史数据
        history_summary = self.operation_history.get_history_summary()
        for item in history_summary:
            time_str = item['timestamp'][:19].replace('T', ' ')
            status = "当前" if item['is_current'] else ("可撤销" if item['can_undo'] else "已完成")
            history_tree.insert("", "end", values=(time_str, item['description'], status))
        
        ttk.Button(button_frame, text="撤销", command=lambda: self._history_undo_and_refresh(history_dialog)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="重做", command=lambda: self._history_redo_and_refresh(history_dialog)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="清空历史", command=lambda: self._clear_history_and_refresh(history_dialog)).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="关闭", command=history_dialog.destroy).pack(side=tk.RIGHT)
    
    def _history_undo_and_refresh(self, dialog):
        """在历史对话框中撤销操作并刷新"""
        self._execute_history_action("undo", "操作已撤销", "无法撤销当前操作", dialog)
     
    def update_history_buttons(self):
        """更新历史操作按钮的状态"""
        # 更新撤销按钮状态
        if hasattr(self, 'undo_button'):
            if self.operation_history.can_undo():
                self.undo_button.configure(state="normal")
            else:
                self.undo_button.configure(state="disabled")
        
        # 更新重做按钮状态
        if hasattr(self, 'redo_button'):
            if self.operation_history.can_redo():
                self.redo_button.configure(state="normal")
            else:
                self.redo_button.configure(state="disabled")
    
    def _history_redo_and_refresh(self, dialog):
        """在历史对话框中重做操作并刷新"""
        self._execute_history_action("redo", "操作已重做", "无法重做操作", dialog)
    
    def _clear_history_and_refresh(self, dialog):
        """清空历史记录并刷新"""
        if messagebox.askyesno("确认", "确定要清空所有操作历史吗？此操作不可撤销。"):
            self.operation_history.clear_history()
            self._show_info_message("成功", "操作历史已清空")
            dialog.destroy()
            self.show_operation_history()  # 重新显示更新后的历史
    
    def show_help_info(self):
        """显示帮助信息对话框"""
        help_dialog = tk.Toplevel(self.root)
        help_dialog.title("帮助信息")
        help_dialog.geometry("800x600")
        
        # 设置对话框位置
        help_dialog.update_idletasks()
        screen_width = help_dialog.winfo_screenwidth()
        screen_height = help_dialog.winfo_screenheight()
        dialog_width = 800
        dialog_height = 600
        x = (screen_width - dialog_width) // 2
        y = (screen_height - dialog_height) // 2
        help_dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")
        
        help_dialog.transient(self.root)
        help_dialog.grab_set()
        
        main_frame = ttk.Frame(help_dialog, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(main_frame, text=f"{APP_NAME} - 帮助信息", 
                               font=("TkDefaultFont", 14, "bold"))
        title_label.pack(pady=(0, 15))
        
        # 创建滚动文本框
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True)
        
        help_text = tk.Text(text_frame, wrap=tk.WORD, font=("TkDefaultFont", 10))
        scrollbar = ttk.Scrollbar(text_frame, orient="vertical", command=help_text.yview)
        help_text.configure(yscrollcommand=scrollbar.set)
        
        help_content = f"""【软件功能概述】
本工具是一个功能强大的目录和文件管理工具，主要用于目录结构复制、文件批量重命名、多格式导出操作和操作历史管理。支持多种复制模式、智能重命名、10种导出格式和完善的撤销/重做功能。

【主要功能模块】

1. 【目录选择与浏览】
   • 源目录选择：选择要操作的源目录，支持中文路径
   • 目标目录选择：选择复制或导出的目标位置
   • 目录树浏览：以树形结构显示目录和文件，支持展开/折叠
   • 复选框选择：可以选择性地操作特定的目录或文件
   • 文件大小显示：实时显示文件和目录大小（KB单位）
   • 类型标识：清晰区分文件和目录类型

2. 【操作模式】（5种模式）
   • 复制单层目录：仅复制源目录的第一层空目录结构
   • 复制选中层级目录：复制在树形视图中勾选的目录结构
   • 复制所有层级目录：递归复制整个目录树结构（空目录）
   • 复制选定目录和文件：灵活选择要复制的文件和目录，支持混合选择
   • 导出目录和文件的名称：将目录结构导出为10种文档格式

3. 【显示筛选】（3种模式）
   • 默认显示：同时显示目录和文件
   • 只显示目录：仅在列表中显示目录，便于查看目录层次
   • 只显示文件：仅在列表中显示文件，便于查看文件清单

4. 【多格式导出功能】（10种格式）
   • 导出目录文件名称：生成目录结构文档，支持10种不同格式
   
   格式详细说明：
   ① TXT格式：纯文本树形结构，简洁清晰，适合快速查看
   ② HTML格式：静态网页格式，美观易读
   ③ HTML格式(含链接)：在HTML基础上添加可点击链接，直接打开文件
   ④ Markdown格式：适合文档编写和版本控制，支持GitHub等平台
   ⑤ Markdown(含链接)：使用Markdown列表输出可点击文件链接
   ⑥ DOCX格式：Microsoft Word文档格式，专业文档输出
   ⑦ DOCX格式(含链接)：Word文档格式，支持真正的超链接，可直接点击打开文件
   ⑧ JSON格式：标准JSON树形结构，适合程序化处理
   ⑨ XLSX格式：Excel表格格式，行列展示目录结构
   ⑩ XLSX格式(含链接)：Excel表格格式，支持可点击的文件链接
   
   • 自动命名：基于当前时间戳和源目录名生成文件名（格式：时间戳_目录名_目录结构.扩展名）
   • 结构化输出：包含完整的层级关系和文件信息
   • 链接功能：支持file://协议的本地文件链接，兼容中文路径
   • 格式化支持：DOCX格式支持居中标题、右对齐时间戳、分层缩进

5. 【批量重命名功能】
   • 字符替换重命名：
     - 重命名本级目录：当前级别目录的字符串查找替换
     - 重命名全部：递归重命名所有层级的文件和目录
   
   • 多维重命名：
     - 本级目录重命名：先在目录树中勾选本级目录（可"全选/取消全选"），再设置前缀、后缀、序号编号，仅对已勾选目录改名
     - 本级文件重命名：先在目录树中勾选本级文件（可"全选/取消全选"），保持文件扩展名，支持批量编号
     - 序号格式：数字（001,002...）、字母（A,B,C...或a,b,c...）、罗马数字（I,II,III...）
     - 连接符设置：自定义前缀、后缀与序号间的连接符
   
   • 名称修改（高级重命名）：
     - 全层级目录名称修改：先勾选各层级的目录（可逐个勾选、全选或取消全选），再执行精确匹配、通配符（直观替换）、正则表达式改名
     - 各层级文件名称修改：先勾选各层级的文件（可逐个勾选、全选或取消全选），再执行精确匹配、通配符（直观替换）、正则表达式改名
     - 预览功能：重命名前可预览结果
     - 条件替换：复杂的重命名规则
     - 安全检查：防止重名冲突

6. 【操作历史管理】
   • 智能历史记录：自动记录所有重命名和复制操作
   • 撤销功能：可以撤销最近的操作，支持多级撤销
   • 重做功能：可以重做已撤销的操作
   • 操作历史查看：详细显示操作时间、类型、状态
   • 历史清空：一键清除所有操作历史记录
   • 历史记录：当前版本主要记录复制和重命名操作的历史信息
   • 状态管理：实时显示可撤销/重做状态

7. 【选择操作】
   • 全选：选择当前目录下的所有项目
   • 取消全选：取消所有选择
   • 单项选择：通过复选框选择特定项目
   • 智能选择：根据文件类型或大小进行筛选

【操作流程】

1. 基本复制流程：
   ① 选择源目录（点击"浏览"按钮）
   ② 选择目标目录（点击"浏览"按钮）
   ③ 选择操作模式（复制或导出）
   ④ 在目录树中选择要操作的项目（可选）
   ⑤ 点击"开始复制"按钮执行操作
   ⑥ 查看操作结果和完成提示

2. 多格式导出流程：
   ① 选择源目录（点击"浏览"按钮）
   ② 选择目标目录（导出文件保存位置）
   ③ 选择"导出目录和文件的名称"模式
   ④ 选择导出格式（10种格式可选）：
      • TXT格式：纯文本树形结构
      • HTML格式：静态网页格式
      • HTML格式(含链接)：可点击链接直接打开文件
      • Markdown格式：适合文档编写
      • Markdown(含链接)：使用Markdown列表输出文件链接
      • DOCX格式：Word文档格式
      • DOCX格式(含链接)：支持真正超链接的Word文档
      • JSON格式：标准JSON树形结构，适合程序化处理
      • XLSX格式：Excel表格格式，行列展示目录结构
      • XLSX格式(含链接)：Excel表格格式，支持可点击的文件链接
   ⑤ 点击"开始复制"按钮执行导出
   ⑥ 系统自动生成并保存导出文件到目标目录
   ⑦ 查看导出成功提示和文件位置

3. 重命名流程：
   ① 选择源目录
   ② 在目录树中浏览到要重命名的目录
   ③ 选择相应的重命名功能
   ④ 设置重命名参数（前缀、后缀、序号等）
   ⑤ 预览重命名结果（高级重命名）
   ⑥ 执行重命名操作
   ⑦ 查看操作历史记录

【注意事项】

• 重命名操作会直接修改文件系统，建议先备份重要数据
• 操作历史功能可用于撤销已记录的复制和重命名操作
• 复制大量文件时请耐心等待，系统会显示处理进度
• 目标目录如果已存在同名文件，系统会智能处理冲突
• 导出功能支持10种格式：TXT、HTML、HTML(含链接)、Markdown、Markdown(含链接)、DOCX、DOCX(含链接)、JSON、XLSX、XLSX(含链接)
• DOCX格式需要python-docx库支持，如未安装会自动提示安装命令
• 链接版本的导出格式支持直接点击打开文件，适合制作可交互的目录文档
• 高级重命名支持预览功能，建议先预览再执行
• 程序会自动处理中文路径和特殊字符，支持Unicode文件名
• 操作过程中如遇到权限问题，会给出明确提示
• DOCX格式的超链接功能使用file://协议，兼容Windows资源管理器

【快捷操作技巧】

• 双击目录项可以快速展开/折叠
• 使用"全选"/"取消全选"可以快速选择项目
• 目录树采用懒加载：未展开的目录不会预先加载，"全选"仅覆盖已加载的项目；如需选择全部文件，请先展开所有目录，或切换到"只显示文件"筛选后再"全选"
• "撤销"/"重做"按钮可以快速恢复操作
• "操作历史"可以查看详细的操作记录和状态
• 树形视图支持滚动浏览大型目录结构
• 复选框状态会实时反映选择情况

【错误处理与故障排除】

如果在使用过程中遇到问题，请检查：
1. 文件路径是否正确且存在
2. 是否有足够的磁盘空间
3. 是否有相应的文件操作权限
4. 目标目录是否可写
5. 源目录是否被其他程序占用
6. 文件名是否包含非法字符
7. 网络驱动器连接是否正常

常见问题解决：
• 权限不足：以管理员身份运行程序
• 路径过长：使用较短的目标路径
• 文件被占用：关闭相关程序后重试
• 操作失败：查看操作历史了解详细错误信息

【技术特性】

• 异常处理：完善的错误捕获和用户友好提示
• 内存优化：高效处理大型目录结构
• 安全机制：操作前验证，关键操作具备历史记录支持
• 界面友好：直观的操作流程和状态反馈
• 跨平台：基于Python tkinter，支持Windows系统
• 多格式支持：10种导出格式，满足不同使用场景
• 超链接技术：DOCX格式使用底层XML API实现真正的可点击超链接
• 文件URI转换：支持本地文件路径转换为file://协议，兼容中文路径
• 依赖管理：自动检测python-docx库，缺失时提供安装指导
• 代码隔离：新功能完全独立，不影响原有功能稳定性

【版本信息】
当前版本：{APP_VERSION}
开发日期：{APP_DEVELOP_DATE}
更新日期：{APP_RELEASE_DATE}
作者：{APP_AUTHOR}

【当前版本主要能力】
• 全层级目录名称修改功能：可先勾选部分或全部目录（各层级），再执行精确匹配、通配符（直观替换）、正则表达式改名
• 各层级文件名称修改功能：可先勾选部分或全部文件（各层级），再执行精确匹配、通配符（直观替换）、正则表达式改名
• 多维重命名本级目录名/文件名：仅对已勾选的本级目录/文件执行前缀、后缀、序号编号改名
• 10种导出格式支持：TXT、HTML、HTML(含链接)、Markdown、Markdown(含链接)、DOCX、DOCX(含链接)、JSON、XLSX、XLSX(含链接)
• DOCX格式真正超链接功能：使用底层XML API实现可点击的文件链接
• 文件URI转换技术：支持file://协议，兼容中文路径
• UI界面优化：导出格式选项单行紧凑布局
• 依赖管理增强：自动检测python-docx、XlsxWriter库并提供安装指导
• 代码隔离设计：新功能完全独立，确保原有功能稳定性

【历史版本补充】
V1.5 功能：
• 多维重命名本级目录名/文件名改为先勾选再重命名，支持"全选/取消全选"
• 全层级目录名称修改：先勾选各层级目录再执行精确/通配符/正则改名
• 各层级文件名称修改：先勾选各层级文件再执行精确/通配符/正则改名
• 新增"关于"对话框，显示版本、作者、协议等信息
V1.4 功能：
• 文件名称修改功能：可先勾选部分或全部文件，再执行精确匹配、通配符（直观替换）、正则表达式改名
• 目录名称修改对话框采用通配符直观替换语义
V1.3 功能：
• 目录大小缓存：刷新树时统一构建缓存，减少大目录重复扫描
• UI更新节流：长循环中轻量刷新界面，降低主线程阻塞风险
V1.1 功能：
• 完善的操作历史管理系统
• 智能撤销/重做功能
• 高级重命名预览功能
• 优化的用户界面和交互体验
• 增强的错误处理和安全机制"""
        
        help_text.insert(tk.END, help_content)
        help_text.configure(state="disabled")  # 设置为只读
        
        help_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 关闭按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(15, 0))
        
        ttk.Button(button_frame, text="关闭", command=help_dialog.destroy).pack(side=tk.RIGHT)

    def show_about_info(self):
        """显示关于对话框（20260906 192644 新增）"""
        about_dialog = tk.Toplevel(self.root)
        about_dialog.title("关于")
        about_dialog.geometry("460x360")
        about_dialog.resizable(False, False)

        # 设置对话框位置：水平居中，垂直居中
        about_dialog.update_idletasks()
        screen_width = about_dialog.winfo_screenwidth()
        screen_height = about_dialog.winfo_screenheight()
        dialog_width = 460
        dialog_height = 360
        x = (screen_width - dialog_width) // 2
        y = (screen_height - dialog_height) // 2
        about_dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

        about_dialog.transient(self.root)
        about_dialog.grab_set()

        main_frame = ttk.Frame(about_dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 软件名称
        ttk.Label(main_frame, text=APP_NAME, font=("TkDefaultFont", 15, "bold"),
                  foreground="#2B6CB0").pack(pady=(10, 5))

        # 分隔线
        ttk.Separator(main_frame, orient='horizontal').pack(fill=tk.X, pady=10)

        # 版本与作者信息
        info_frame = ttk.Frame(main_frame)
        info_frame.pack(fill=tk.X)
        info_lines = [
            f"版    本：{APP_VERSION}",
            f"构建日期：{APP_BUILD_DATE}",
            f"发布日期：{APP_RELEASE_DATE}",
            f"开 发 者：{APP_AUTHOR}",
            f"可执行文件：{APP_EXECUTABLE_NAME}.exe",
        ]
        for line in info_lines:
            ttk.Label(info_frame, text=line, font=("TkDefaultFont", 10)).pack(anchor='w', pady=2)

        # 协议与简介
        ttk.Separator(main_frame, orient='horizontal').pack(fill=tk.X, pady=(12, 8))
        ttk.Label(main_frame, text="开源协议：MIT License",
                  font=("TkDefaultFont", 9), foreground="#666666").pack(anchor='w')
        ttk.Label(main_frame, text="运行平台：Windows 7/8/10/11",
                  font=("TkDefaultFont", 9), foreground="#666666").pack(anchor='w')
        try:
            app_path = os.path.dirname(os.path.abspath(sys.executable if getattr(sys, 'frozen', False) else __file__))
        except Exception:
            app_path = ""
        ttk.Label(main_frame, text=f"程序路径：{app_path}",
                  font=("TkDefaultFont", 8), foreground="#888888").pack(anchor='w', pady=(6, 0))

        # 关闭按钮
        ttk.Button(main_frame, text="确定", command=about_dialog.destroy).pack(pady=(15, 0))

    def create_main_interface(self):
        outer_frame = tk.Frame(self.root, bd=3, relief='ridge')  # 20260402 113200 双线边框宽度增加
        outer_frame.grid(row=0, column=0, sticky='nsew')
        
        # 主框架
        main_frame = ttk.Frame(outer_frame, padding="10")  # 20260402 112700 增大主界面四边内边距以提升留白与协调性
        main_frame.grid(row=0, column=0, sticky='nsew')
        
        # 配置根窗口和主框架的网格权重
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)
        outer_frame.grid_rowconfigure(0, weight=1)  # 20260402 113000 外层边框容器随窗口缩放
        outer_frame.grid_columnconfigure(0, weight=1)  # 20260402 113000 外层边框容器随窗口缩放
        main_frame.grid_columnconfigure(1, weight=1)  # 让第1列（输入框列）可扩展
        
        # 设置样式
        style = ttk.Style()
        try:
            style.theme_use('vista')  # 20260402 101800 使用Windows更协调的主题
        except Exception:
            pass

        primary_color = "#2B6CB0"  # 20260402 101800 统一界面主色
        danger_color = "#C53030"  # 20260402 101800 危险动作高亮色
        muted_color = "#666666"  # 20260402 101800 次要信息颜色
        pad_x = 6  # 20260402 101800 统一界面间距
        pad_y = 4  # 20260402 101800 统一界面间距

        style.configure('Treeview', rowheight=30, font=('TkDefaultFont', 11))  # 20260402 101800 列表框保持原字号
        style.configure('Treeview.Heading', font=('TkDefaultFont', 10), relief='ridge', borderwidth=1)  # 20260402 113400 表头分隔更明显
        style.configure('Treeview', background='white', foreground='black', relief='solid', borderwidth=1)  # 20260402 113400 表格边框更明显
        style.map('Treeview', background=[('selected', '#E8F1FF')])
        style.layout('Treeview.Item', [
            ('Treeitem.padding', {'sticky': 'nswe', 'children': [
                ('Treeitem.indicator', {'side': 'left', 'sticky': ''}),
                ('Treeitem.image', {'side': 'left', 'sticky': ''}),
                ('Treeitem.text', {'side': 'left', 'sticky': ''})
            ]})
        ])
        
        # 软件标题
        title_label = tk.Label(main_frame, text=APP_NAME, 
                              font=self.title_font, fg=primary_color)  # 20260402 101800 统一界面主色
        title_label.grid(row=0, column=0, columnspan=4, pady=(0, 10))
        
        # 源目录选择
        source_label = ttk.Label(main_frame, text="请先选择源目录:", font=self.bold_label_font)
        source_label.grid(row=1, column=0, sticky='w')
        self.source_entry = ttk.Entry(main_frame, textvariable=self.source_dir)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.source_entry.grid(row=1, column=1, padx=pad_x, pady=pad_y, sticky='ew')  # 20260402 101800 统一界面间距
        self.source_browse_button = ttk.Button(main_frame, text="浏览", command=self.select_source)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.source_browse_button.grid(row=1, column=2, padx=pad_x, pady=pad_y)  # 20260402 101800 统一界面间距
        
        # 目标目录选择
        dest_label = ttk.Label(main_frame, text="请选择目标目录:", font=self.bold_label_font)
        dest_label.grid(row=2, column=0, sticky='w', pady=pad_y)  # 20260402 101800 统一界面间距
        self.dest_entry = ttk.Entry(main_frame, textvariable=self.dest_dir)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.dest_entry.grid(row=2, column=1, padx=pad_x, pady=pad_y, sticky='ew')  # 20260402 101800 统一界面间距
        self.dest_browse_button = ttk.Button(main_frame, text="浏览", command=self.select_dest)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.dest_browse_button.grid(row=2, column=2, padx=pad_x, pady=pad_y)  # 20260402 101800 统一界面间距
        
        # 复制模式选择
        mode_frame = ttk.LabelFrame(main_frame, text="复制目录结构和导出文件名功能：", padding="5")
        mode_frame.grid(row=3, column=0, columnspan=3, sticky='ew', pady=10)
        
        # 设置LabelFrame标题颜色为蓝色，加粗
        style.configure('Blue.TLabelframe.Label', foreground=primary_color, font=('TkDefaultFont', 10, 'bold'))  # 20260402 101800 统一界面主色与字号
        mode_frame.configure(style='Blue.TLabelframe')
        
        self.mode_radio_buttons = []  # 20260402 100551 记录模式单选按钮，任务执行期间可禁用
        rb = ttk.Radiobutton(mode_frame, text="仅复制一层目录", value="single_level", variable=self.copy_mode, command=self.on_mode_change)
        rb.grid(row=0, column=0, padx=(5,2))
        self.mode_radio_buttons.append(rb)
        rb = ttk.Radiobutton(mode_frame, text="复制选定层级目录", value="selected_levels", variable=self.copy_mode, command=self.on_mode_change)
        rb.grid(row=0, column=1, padx=2)
        self.mode_radio_buttons.append(rb)
        rb = ttk.Radiobutton(mode_frame, text="复制所有层级目录", value="all_levels", variable=self.copy_mode, command=self.on_mode_change)
        rb.grid(row=0, column=2, padx=2)
        self.mode_radio_buttons.append(rb)
        rb = ttk.Radiobutton(mode_frame, text="复制选定目录和文件", value="custom", variable=self.copy_mode, command=self.on_mode_change)
        rb.grid(row=0, column=3, padx=2)
        self.mode_radio_buttons.append(rb)
        rb = ttk.Radiobutton(mode_frame, text="导出目录和文件的名称", value="export_names", variable=self.copy_mode, command=self.on_mode_change)
        rb.grid(row=0, column=4, padx=2)
        self.mode_radio_buttons.append(rb)
        ttk.Label(mode_frame, text="显示筛选:").grid(row=1, column=0, sticky='w', padx=(5, 2), pady=(6, 0))
        self.display_radio_buttons = []  # 20260402 100551 记录显示筛选单选按钮，任务执行期间可禁用
        rb = ttk.Radiobutton(mode_frame, text="默认显示", value="normal", variable=self.tree_display_mode, command=self.on_tree_display_mode_change)  # 20260402 101800 统一控件风格（ttk）
        rb.grid(row=1, column=1, padx=2, pady=(6, 0), sticky='w')
        self.display_radio_buttons.append(rb)
        rb = ttk.Radiobutton(mode_frame, text="只显示目录", value="directories_only", variable=self.tree_display_mode, command=self.on_tree_display_mode_change)  # 20260402 101800 统一控件风格（ttk）
        rb.grid(row=1, column=2, padx=2, pady=(6, 0), sticky='w')
        self.display_radio_buttons.append(rb)
        rb = ttk.Radiobutton(mode_frame, text="只显示文件", value="files_only", variable=self.tree_display_mode, command=self.on_tree_display_mode_change)  # 20260402 101800 统一控件风格（ttk）
        rb.grid(row=1, column=3, padx=2, pady=(6, 0), sticky='w')
        self.display_radio_buttons.append(rb)
        self.directory_size_checkbutton = ttk.Checkbutton(mode_frame, text="显示目录大小", variable=self.show_directory_sizes, command=self.on_directory_size_toggle)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.directory_size_checkbutton.grid(row=1, column=4, padx=8, pady=(6, 0), sticky='w')
        
        # 导出格式选择框架
        self.export_format_frame = ttk.LabelFrame(main_frame, text="导出格式", padding="5")
        self.export_format_frame.grid(row=4, column=0, columnspan=3, sticky='ew', pady=5)
        
        # 将所有导出格式选项放在一行，缩小间距
        ttk.Radiobutton(self.export_format_frame, text="TXT", value="txt",  # 20260402 111300 导出格式标签去掉“格式”
                       variable=self.export_format).grid(row=0, column=0, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="HTML", value="html",  # 20260402 111300 导出格式标签去掉“格式”
                       variable=self.export_format).grid(row=0, column=1, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="HTML(含链接)", value="html_link",  # 20260402 111300 导出格式标签去掉“格式”
                       variable=self.export_format).grid(row=0, column=2, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="Markdown", value="md",  # 20260402 111300 导出格式标签去掉“格式”
                       variable=self.export_format).grid(row=0, column=3, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="Markdown(含链接)", value="md_link", 
                       variable=self.export_format).grid(row=0, column=4, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="DOCX", value="docx",  # 20260402 111300 导出格式标签去掉“格式”
                       variable=self.export_format).grid(row=0, column=5, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="DOCX(含链接)", value="docx_link",  # 20260402 111300 导出格式标签去掉“格式”
                       variable=self.export_format).grid(row=0, column=6, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="JSON", value="json",  # 20260402 111300 标签去掉'.'并转大写
                       variable=self.export_format).grid(row=0, column=7, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="XLSX", value="xlsx",  # 20260402 111300 标签去掉'.'并转大写
                       variable=self.export_format).grid(row=0, column=8, padx=5)
        ttk.Radiobutton(self.export_format_frame, text="XLSX(含链接)", value="xlsx_link",  # 20260402 111300 标签去掉'.'并转大写
                       variable=self.export_format).grid(row=0, column=9, padx=5)
        
        # 初始隐藏导出格式选择
        self.export_format_frame.grid_remove()
        # 20260906 205815 新增：应用用户上次选择的导出格式
        self._apply_saved_export_format()
        
        # 重命名按钮框架
        rename_frame = ttk.Frame(main_frame)
        rename_frame.grid(row=5, column=0, columnspan=3, sticky='ew', pady=5)
        
        # 添加重命名功能标签（与框架标题字体大小一致，加粗，蓝色）
        rename_label = tk.Label(rename_frame, text="重命名功能:", 
                               font=self.bold_label_font, fg=primary_color)  # 20260402 101800 统一界面主色与字号
        rename_label.pack(side=tk.LEFT, padx=(0, 10))
        
        self.rename_buttons = []  # 20260402 100551 记录重命名按钮，任务执行期间可禁用
        btn = ttk.Button(rename_frame, text="本级目录字符替换", command=self.rename_current_level)
        btn.pack(side=tk.LEFT, padx=(0,5))
        self.rename_buttons.append(btn)
        btn = ttk.Button(rename_frame, text="全部目录字符替换", command=self.rename_all_items)
        btn.pack(side=tk.LEFT, padx=5)
        self.rename_buttons.append(btn)
        btn = ttk.Button(rename_frame, text="多维重命名本级目录名", command=self.multi_rename_current_level)
        btn.pack(side=tk.LEFT, padx=5)
        self.rename_buttons.append(btn)
        btn = ttk.Button(rename_frame, text="多维重命名本级文件名", command=self.multi_rename_current_files)
        btn.pack(side=tk.LEFT, padx=5)
        self.rename_buttons.append(btn)
        btn = ttk.Button(rename_frame, text="全层级目录名称修改", command=self.advanced_rename_directories)  # 20260906 184352 功能名由“全部目录名修改”改为“全层级目录名称修改”，支持先勾选部分目录再修改
        btn.pack(side=tk.LEFT, padx=5)
        self.rename_buttons.append(btn)
        btn = ttk.Button(rename_frame, text="各层级文件名称修改", command=self.advanced_rename_files)  # 20260906 184352 功能名由“文件名称修改”改为“各层级文件名称修改”（原20260828 104445 由“全部文件名修改”改为“文件名称修改”）
        btn.pack(side=tk.LEFT, padx=5)
        self.rename_buttons.append(btn)
        
        # 文件树视图
        tree_frame = ttk.Frame(main_frame)
        tree_frame.grid(row=6, column=0, columnspan=4, sticky='nsew', pady=5)
        
        # 修改后的树形视图配置
        self.tree = ttk.Treeview(tree_frame, selectmode='none', height=6)
        self.tree["columns"] = ("checked", "type", "size")
        # 20260906 201500 修改：列宽改为支持拖动调整并记忆，全部列固定宽度（stretch=False），
        # 确保“显示宽度=配置宽度”，拖动所见即所得，保存/恢复的列宽与用户看到的完全一致
        self.tree.column("#0", width=270, stretch=False)
        self.tree.column("checked", width=60, anchor='center', stretch=False)
        self.tree.column("type", width=80, stretch=False)
        self.tree.column("size", width=80, anchor='e', stretch=False)
        
        # 配置树形视图的标签样式，添加行分隔效果
        self.tree.tag_configure('oddrow', background='#F2F6FF')  # 20260402 113400 行间底色更接近表格分隔视觉
        self.tree.tag_configure('evenrow', background='white')
        
        self.tree.heading("#0", text="名称")
        self.tree.heading("checked", text="选择")
        self.tree.heading("type", text="类型")
        self.tree.heading("size", text="大小")  # 20260402 101800 简化列标题更协调
        
        # 添加点击事件绑定
        self.tree.bind('<ButtonRelease-1>', self.on_tree_click)
        self.tree.bind('<<TreeviewOpen>>', self.on_tree_open)  # 20260402 085901 目录树懒加载：展开时再加载子节点
        # 20260906 201500 新增：目录树表头列宽拖动调整（在列分隔线上按住左键拖动即可调整列宽）
        self.tree.bind('<ButtonPress-1>', self._on_tree_button_press)
        self.tree.bind('<B1-Motion>', self._on_tree_button_motion)
        # 20260906 201500 新增：应用用户上次自定义的目录树列宽
        self._apply_saved_tree_column_widths()
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        h_scrollbar = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

        # 20260906 205815 新增：目录树搜索条（在树上方）
        search_bar = ttk.Frame(tree_frame)
        search_bar.grid(row=0, column=0, columnspan=2, sticky='ew', pady=(0, 4))
        ttk.Label(search_bar, text="搜索：").pack(side=tk.LEFT)
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_bar, textvariable=self.search_var)
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(2, 6))
        search_entry.bind('<Return>', lambda e: self._search_in_tree())
        ttk.Button(search_bar, text="搜索", command=self._search_in_tree).pack(side=tk.LEFT)
        ttk.Button(search_bar, text="清除", command=lambda: (self.search_var.set(""),)).pack(side=tk.LEFT, padx=(4, 0))
        self.search_result_var = tk.StringVar()
        ttk.Label(search_bar, textvariable=self.search_result_var, foreground="#666666").pack(side=tk.LEFT, padx=(8, 0))

        # 布局树视图和滚动条
        self.tree.grid(row=1, column=0, sticky='nsew')
        v_scrollbar.grid(row=1, column=1, sticky='ns')
        h_scrollbar.grid(row=2, column=0, sticky='ew')
        tree_frame.grid_rowconfigure(1, weight=1)  # 20260906 205815 让目录树占据剩余空间
        

        
        # 操作按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=8, column=0, columnspan=4, pady=10)
        
        # 操作按钮（单行布局）
        main_button_frame = ttk.Frame(button_frame)
        main_button_frame.pack(fill=tk.X, pady=(0, 5))
        
        # 回退和重做按钮（放在最左侧）
        # 配置红色文字样式
        style.configure('Red.TButton', foreground=danger_color)  # 20260402 101800 统一危险动作颜色
        
        self.undo_button = ttk.Button(main_button_frame, text="撤销 ↶", command=self.undo_last_operation, state="disabled", style='Red.TButton')
        self.undo_button.pack(side=tk.LEFT, padx=5)
        
        self.redo_button = ttk.Button(main_button_frame, text="重做 ↷", command=self.redo_last_operation, state="disabled", style='Red.TButton')
        self.redo_button.pack(side=tk.LEFT, padx=5)
        
        # 操作历史按钮
        self.history_button = ttk.Button(main_button_frame, text="操作历史", command=self.show_operation_history)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.history_button.pack(side=tk.LEFT, padx=5)
        
        # 选择按钮
        self.select_all_button = ttk.Button(main_button_frame, text="全选", command=self.select_all)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.select_all_button.pack(side=tk.LEFT, padx=5)
        self.deselect_all_button = ttk.Button(main_button_frame, text="取消全选", command=self.deselect_all)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.deselect_all_button.pack(side=tk.LEFT, padx=5)
        
        # 创建开始复制按钮，增大字体和尺寸
        self.start_button = ttk.Button(main_button_frame, text="开始复制或导出", command=self.start_copy)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.start_button.configure(width=16)  # 增加按钮宽度
        # 为按钮设置样式以应用更大字体和蓝色
        style.configure('Large.TButton', font=self.button_font, foreground=primary_color)  # 20260402 101800 统一界面主色
        self.start_button.configure(style='Large.TButton')
        self.start_button.pack(side=tk.LEFT, padx=5, ipady=4)  # 增加按钮高度
        self.cancel_copy_button = ttk.Button(main_button_frame, text="清空选择", command=self.cancel_copy)  # 20260402 102000 调整按钮文案更明确（清空勾选项）
        self.cancel_copy_button.pack(side=tk.LEFT, padx=5)
        self.help_button = ttk.Button(main_button_frame, text="帮助信息", command=self.show_help_info)  # 20260402 100551 记录控件引用，任务执行期间可禁用
        self.help_button.pack(side=tk.LEFT, padx=5)
        self.about_button = ttk.Button(main_button_frame, text="关于", command=self.show_about_info)  # 20260906 192644 新增“关于”按钮，显示软件版本/作者/协议等信息
        self.about_button.pack(side=tk.LEFT, padx=5)
        self.cancel_task_button = ttk.Button(main_button_frame, text="停止当前任务", command=self.cancel_current_task, state="disabled")  # 20260402 111139 将停止任务按钮移至主按钮行，与帮助信息同排
        self.cancel_task_button.pack(side=tk.LEFT, padx=5)

        self.root.update_idletasks()  # 20260402 115200 计算按钮请求宽度后再按比例收窄（仅改宽度不改高度）
        self._shrink_bottom_command_button_widths(0.80)  # 20260402 115400 底部命令按钮宽度缩小到现在的80%（仅改宽度不改高度）
        
        # 更新按钮状态
        self.update_history_buttons()
        
        # 状态栏（进度与任务取消）
        status_frame = ttk.Frame(main_frame)  # 20260402 101800 将进度与停止按钮放入底部状态栏更协调
        status_frame.grid(row=9, column=0, columnspan=4, sticky='ew', pady=(2, 2))
        status_frame.grid_columnconfigure(0, weight=1)
        self.progress_var = tk.StringVar(value="")  # 20260402 091600 进度文本
        self.progress_label = ttk.Label(status_frame, textvariable=self.progress_var, foreground=primary_color)  # 20260402 101800 统一界面主色
        self.progress_label.grid(row=0, column=0, sticky='w')  # 20260402 111139 状态栏仅显示进度文本

        # 添加版本信息和作者信息
        info_frame = ttk.Frame(main_frame)
        info_frame.grid(row=10, column=0, columnspan=4, sticky='ew', pady=5)  # 20260402 101800 状态栏占用一行，版本信息下移
        
        # 左下角版本信息
        version_label = tk.Label(info_frame, text=f"{APP_VERSION}   {APP_RELEASE_DATE}", fg=muted_color, font=('TkDefaultFont', 9))  # 20260402 101800 次要信息使用灰阶
        version_label.pack(side=tk.LEFT)
        
        # 右下角作者信息
        author_label = tk.Label(info_frame, text=f"作者：{APP_AUTHOR}", fg=muted_color, font=('TkDefaultFont', 9))  # 20260402 101800 次要信息使用灰阶
        author_label.pack(side=tk.RIGHT)
        
        # 配置grid权重
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)
        main_frame.grid_columnconfigure(1, weight=1)
        main_frame.grid_rowconfigure(6, weight=1)  # 让文件树视图行可扩展
        tree_frame.grid_columnconfigure(0, weight=1)
        tree_frame.grid_rowconfigure(0, weight=1)  # 让树视图可扩展

    def _shrink_bottom_command_button_widths(self, ratio):  # 20260402 115200 底部命令按钮宽度按比例收窄（仅改宽度）
        buttons = []
        for attr in [
            "undo_button", "redo_button", "history_button",
            "select_all_button", "deselect_all_button",
            "cancel_copy_button", "help_button", "about_button",
        ]:
            if hasattr(self, attr):
                btn = getattr(self, attr)
                if btn:
                    buttons.append(btn)
        for btn in buttons:
            self._shrink_button_width_to_ratio(btn, ratio)

    def _shrink_button_width_to_ratio(self, button, ratio):  # 20260402 115200 将ttk按钮宽度按像素比例换算到字符宽度
        try:
            current_px = int(button.winfo_reqwidth())
            if current_px <= 0:
                return
            desired_px = max(1, int(current_px * ratio))
            try:
                f = font.nametofont(button.cget("font"))
            except Exception:
                f = font.nametofont("TkDefaultFont")
            unit_px = max(1, int(f.measure("0")))
            desired_chars = max(1, int(desired_px / unit_px))
            button.configure(width=desired_chars)
        except Exception:
            pass

    def on_mode_change(self):
        """当复制模式改变时的处理"""
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，暂不支持切换模式。")  # 20260402 100551 任务运行期间禁止切换模式以避免状态混乱
            return
        mode = self.copy_mode.get()
        self.checked_items.clear()  # 清除之前的选择
        
        # 根据模式显示或隐藏导出格式选择
        if mode == "export_names":
            self.export_format_frame.grid()
        else:
            self.export_format_frame.grid_remove()
            
        self.refresh_tree()  # 刷新树形视图

    def on_tree_display_mode_change(self):
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，暂不支持切换显示筛选。")  # 20260402 100551 任务运行期间禁止切换显示筛选以避免状态混乱
            return
        self.refresh_tree()

    def on_directory_size_toggle(self):
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，暂不支持切换目录大小显示。")  # 20260402 100551 任务运行期间禁止切换目录大小显示以避免状态混乱
            return
        self.refresh_tree()  # 20260402 085901 切换目录大小显示后刷新目录树
    
    def _set_task_ui_state(self, running):  # 20260402 100551 任务执行期间禁用关键控件，防止并发操作导致整体崩溃
        state = "disabled" if running else "normal"
        widgets = []
        for attr in [
            "source_entry", "dest_entry", "source_browse_button", "dest_browse_button",
            "start_button", "history_button", "select_all_button", "deselect_all_button",
            "cancel_copy_button", "help_button", "directory_size_checkbutton"
        ]:
            if hasattr(self, attr):
                widget = getattr(self, attr)
                if widget:
                    widgets.append(widget)
        if hasattr(self, "rename_buttons"):
            widgets.extend(list(self.rename_buttons))
        for widget in widgets:
            try:
                widget.configure(state=state)
            except Exception:
                pass
        for group_attr in ["mode_radio_buttons", "display_radio_buttons"]:
            if hasattr(self, group_attr):
                for rb in getattr(self, group_attr) or []:
                    try:
                        rb.configure(state=state)
                    except Exception:
                        pass
        try:
            self.undo_button.configure(state="disabled" if running else self.undo_button.cget("state"))
            self.redo_button.configure(state="disabled" if running else self.redo_button.cget("state"))
        except Exception:
            pass
        if not running:
            try:
                self.update_history_buttons()  # 20260402 100551 恢复后同步撤销/重做按钮真实状态
            except Exception:
                pass
    
    class TaskRunner:  # 20260402 091600 分片任务调度器（可取消）
        def __init__(self, app, description, total=None, step_size=50):
            self.app = app
            self.description = description
            self.total = total
            self.step_size = step_size
            self.queue = []
            self._after_id = None
            self._running = False
            self._on_done = None
            self._step_fn = None
            self._completed = 0
            self._failed = 0
            self._skipped = 0
            self._batch_operations = []
        
        def start(self, items, step_fn, on_done):
            if self._running:
                return False
            self.queue = list(items)
            self._step_fn = step_fn
            self._on_done = on_done
            self._running = True
            self.app.task_running = True  # 20260402 091600 标记运行中
            self.app.cancel_requested = False  # 20260402 091600 清除取消标志
            self.app._set_task_ui_state(True)  # 20260402 100551 任务启动时禁用关键控件
            self._update_progress("开始")
            self._schedule_next()
            return True
        
        def cancel(self):
            self.app.cancel_requested = True  # 20260402 091600 外部取消
        
        def _schedule_next(self):
            if self.app._test_force_sync:  # 20260402 091600 测试模式：同步执行
                self._process_step()
                return
            self._after_id = self.app.root.after(0, self._process_step)
        
        def _process_step(self):
            if self.app.cancel_requested:
                self._finish(cancelled=True)
                return
            count = 0
            while self.queue and count < self.step_size:
                item = self.queue.pop(0)
                try:
                    result = self._step_fn(item, self._batch_operations)
                    status = result
                    new_items = None
                    if isinstance(result, tuple) and len(result) == 2:  # 20260402 094929 支持动态追加队列（用于递归遍历类任务）
                        status, new_items = result
                    elif isinstance(result, dict):  # 20260402 094929 支持字典返回值（扩展兼容）
                        status = result.get("status")
                        new_items = result.get("new_items")
                    if new_items:
                        self.queue.extend(list(new_items))  # 20260402 094929 将新任务追加到队列尾部
                    if status == "skipped":
                        self._skipped += 1
                    elif status == "failed":  # 20260402 100551 支持通过返回状态统计失败（不依赖抛异常）
                        self._failed += 1
                    else:
                        self._completed += 1
                except Exception:
                    self._failed += 1
                count += 1
            self._update_progress("处理中")
            if self.app.cancel_requested:
                self._finish(cancelled=True)
                return
            if not self.queue:
                self._finish(cancelled=False)
            else:
                self._schedule_next()
        
        def _update_progress(self, phase):
            total = self.total if self.total is not None else (self._completed + self._skipped + self._failed + len(self.queue))
            text = f"{self.description}：{phase}，完成 {self._completed}，跳过 {self._skipped}，失败 {self._failed}"
            if total:
                text += f"（剩余 {len(self.queue)}）"
            self.app.progress_var.set(text)  # 20260402 091600 更新进度文本
            self.app.cancel_task_button.configure(state="normal" if self._running else "disabled")
        
        def _finish(self, cancelled):
            self._running = False
            self.app.task_running = False
            self.app.cancel_requested = False
            if self._after_id is not None and not self.app._test_force_sync:
                try:
                    self.app.root.after_cancel(self._after_id)
                except Exception:
                    pass
                self._after_id = None
            if self._on_done:
                try:
                    self._on_done({
                        "completed": self._completed,
                        "skipped": self._skipped,
                        "failed": self._failed,
                        "cancelled": cancelled,
                        "batch_operations": self._batch_operations
                    })
                finally:
                    self._on_done = None
            self.app._set_task_ui_state(False)  # 20260402 100551 任务结束后恢复关键控件
            self.app.progress_var.set("")  # 20260402 091600 清空进度文本
            self.app.cancel_task_button.configure(state="disabled")  # 20260402 091600 禁用取消按钮
    
    def cancel_current_task(self):
        if getattr(self, "_export_collecting", False):  # 20260828 111834 支持取消导出分片收集
            self.cancel_requested = True
            return
        if hasattr(self, 'task_runner') and self.task_runner and self.task_running:  # 20260402 091600 停止当前分片任务
            self.task_runner.cancel()
            
    # 20260906 201500 新增：目录树列宽配置读写与表头列宽拖动
    def _get_config_path(self):
        """获取用户偏好配置文件路径（保存在用户主目录，避免程序目录只读问题）"""
        return os.path.join(os.path.expanduser("~"), CONFIG_FILE_NAME)

    def _load_tree_column_widths(self):
        """读取用户自定义的目录树列宽；无配置或数据异常时返回 None"""
        try:
            path = self._get_config_path()
            if os.path.isfile(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    widths = data.get("tree_column_widths")
                    if isinstance(widths, dict):
                        return widths
        except Exception:
            pass  # 配置读取失败时静默使用默认列宽，不影响软件功能
        return None

    def _save_tree_column_widths(self):
        """保存用户自定义的目录树列宽到用户主目录配置文件"""
        try:
            widths = {}
            for col in TREE_COLUMN_NAMES:
                widths[col] = int(self.tree.column(col, "width"))
            path = self._get_config_path()
            data = {}
            if os.path.isfile(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if not isinstance(data, dict):
                        data = {}
                except Exception:
                    data = {}  # 原配置损坏时重建
            data["tree_column_widths"] = widths
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass  # 配置保存失败时静默忽略，不影响软件功能

    def _apply_saved_tree_column_widths(self):
        """应用用户上次自定义的目录树列宽"""
        widths = self._load_tree_column_widths()
        if not widths:
            return
        for col in TREE_COLUMN_NAMES:
            try:
                w = int(widths.get(col, 0))
            except (TypeError, ValueError):
                continue
            if TREE_COLUMN_MIN_WIDTH <= w <= TREE_COLUMN_MAX_WIDTH:
                self.tree.column(col, width=w)

    def _load_export_format(self):  # 20260906 205815 新增：读取用户上次选择的导出格式
        """读取用户上次选择的导出格式；无配置时返回 None"""
        try:
            path = self._get_config_path()
            if os.path.isfile(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    fmt = data.get("export_format")
                    if isinstance(fmt, str) and fmt:
                        return fmt
        except Exception:
            pass
        return None

    def _save_export_format(self):  # 20260906 205815 新增：保存用户选择的导出格式
        """保存当前导出格式到用户主目录配置文件"""
        try:
            fmt = self.export_format.get()
            path = self._get_config_path()
            data = {}
            if os.path.isfile(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if not isinstance(data, dict):
                        data = {}
                except Exception:
                    data = {}
            data["export_format"] = fmt
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _apply_saved_export_format(self):  # 20260906 205815 新增：应用用户上次选择的导出格式
        """应用用户上次选择的导出格式"""
        fmt = self._load_export_format()
        if fmt:
            self.export_format.set(fmt)

    def _load_window_geometry(self):  # 20260906 223308 新增：读取用户保存的主窗口位置与大小
        """读取用户保存的主窗口几何（WxH+X+Y）；无配置或数据异常时返回 None"""
        try:
            path = self._get_config_path()
            if os.path.isfile(path):
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    geo = data.get("window_geometry")
                    if isinstance(geo, str) and geo:
                        return geo
        except Exception:
            pass  # 配置读取失败时静默使用默认窗口几何，不影响软件功能
        return None

    def _save_window_geometry(self):  # 20260906 223308 新增：保存主窗口位置与大小
        """保存主窗口当前几何（WxH+X+Y）到用户主目录配置文件"""
        try:
            w = self.root.winfo_width()
            h = self.root.winfo_height()
            if w < 1 or h < 1:
                return
            # 以当前实际大小与位置（含左上角偏移，支持多显示器负坐标）记录
            geo = f"{w}x{h}+{self.root.winfo_x()}+{self.root.winfo_y()}"
            path = self._get_config_path()
            data = {}
            if os.path.isfile(path):
                try:
                    with open(path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    if not isinstance(data, dict):
                        data = {}
                except Exception:
                    data = {}
            data["window_geometry"] = geo
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception:
            pass  # 配置保存失败时静默忽略，不影响软件功能

    def _apply_saved_window_geometry(self):  # 20260906 223308 新增：应用用户保存的主窗口位置与大小
        """应用用户保存的主窗口几何；数据非法、小于最小尺寸或窗口主体移出屏幕时回退默认"""
        geo = self._load_window_geometry()
        if not geo:
            return
        mm = re.fullmatch(r"(\d+)x(\d+)([+-]\d+)([+-]\d+)", geo)
        if not mm:
            return
        try:
            w, h, x, y = (int(v) for v in mm.groups())
        except ValueError:
            return
        if w < self.root.minsize()[0] or h < self.root.minsize()[1]:
            return
        # 分辨率/显示器变化容错：窗口中心已移出可见屏幕时，丢弃位置并重新居中，保留用户设定的尺寸
        try:
            sw = self.root.winfo_screenwidth()
            sh = self.root.winfo_screenheight()
        except Exception:
            return
        if not (0 <= x + w // 2 < sw and 0 <= y + h // 2 < sh):
            x = max(0, (sw - w) // 2)
            y = max(0, (sh - h) // 2)
        try:
            self.root.geometry(f"{w}x{h}+{x}+{y}")
        except Exception:
            pass  # 应用失败时保持默认窗口几何，不影响软件启动

    def _on_main_window_close(self):  # 20260906 223308 新增：主窗口关闭前保存窗口位置与大小
        """主窗口关闭回调：先记忆窗口几何，再销毁窗口"""
        self._save_window_geometry()
        self.root.destroy()

    def _on_tree_button_press(self, event):
        """目录树表头列分隔线上按下左键：进入列宽拖动状态"""
        self._tree_resizing = False
        self._tree_resize_col = None
        try:
            region = self.tree.identify("region", event.x, event.y)
        except Exception:
            return
        if region != "separator":
            return
        # 20260906 201500 分隔线归属其左侧列；identify_column 返回位置编号（#0/#1/#2/#3），需映射为列名
        col = self.tree.identify_column(event.x)
        col_map = {"#0": "#0", "#1": "checked", "#2": "type", "#3": "size"}
        target_col = col_map.get(col)
        if target_col is None or target_col not in TREE_COLUMN_NAMES:
            return
        self._tree_resizing = True
        self._tree_resize_col = target_col
        self._tree_resize_start_x = event.x
        self._tree_resize_start_w = int(self.tree.column(target_col, "width"))
        return "break"  # 阻止 Treeview 内部对按下的默认处理

    def _on_tree_button_motion(self, event):
        """按住左键拖动列分隔线：实时调整列宽"""
        if not self._tree_resizing or not self._tree_resize_col:
            return
        new_w = self._tree_resize_start_w + (event.x - self._tree_resize_start_x)
        new_w = max(TREE_COLUMN_MIN_WIDTH, min(TREE_COLUMN_MAX_WIDTH, new_w))
        try:
            self.tree.column(self._tree_resize_col, width=new_w)
        except Exception:
            pass
        return "break"

    # 修改后的树形视图点击事件处理
    def on_tree_click(self, event):
        """处理树形视图的点击事件"""
        # 20260906 201500 新增：列宽拖动结束时跳过复选框处理，并保存用户列宽设置
        if self._tree_resizing:
            self._tree_resizing = False
            self._tree_resize_col = None
            self._save_tree_column_widths()
            return
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            return
            
        item = self.tree.identify_row(event.y)
        if not item:
            return
            
        column = self.tree.identify_column(event.x)
        
        if column == "#1":  # 修正：使用 #1 代替 #2
            current_state = self.tree.set(item, "checked")
            new_state = "🔲" if current_state == "✅" else "✅"
            self.tree.set(item, "checked", new_state)
            
            if new_state == "✅":
                self.checked_items.add(item)
            else:
                self.checked_items.discard(item)
                
            # 强制更新显示
            self.tree.update()
                
    def select_all(self):
        """全选"""
        # 20260906 201500 新增：非勾选模式下点击“全选”时提示用户切换操作模式
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            self._warn_check_mode_required()
            return
        for item in self.get_all_items():
            self.tree.set(item, "checked", "✅")
            self.checked_items.add(item)
        # 强制更新显示
        self.tree.update()

    def deselect_all(self):
        """取消全选"""
        # 20260906 201500 新增：非勾选模式下点击“取消全选”时提示用户切换操作模式（与“全选”行为保持一致）
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            self._warn_check_mode_required()
            return
        for item in self.get_all_items():
            self.tree.set(item, "checked", "🔲")
        self.checked_items.clear()
        # 强制更新显示
        self.tree.update()
            
    def get_all_items(self, parent=''):
        """获取所有项目"""
        items = []
        for item in self.tree.get_children(parent):
            items.append(item)
            items.extend(self.get_all_items(item))
        return items
        
    def select_source(self):
        directory = filedialog.askdirectory(title="选择源目录")
        if directory:
            # 路径规范化处理，增强安全性
            normalized_path = os.path.normpath(os.path.abspath(directory))
            self.source_dir.set(normalized_path)
            self.refresh_tree()
            
    def select_dest(self):
        directory = filedialog.askdirectory(title="选择目标目录")
        if directory:
            # 路径规范化处理，增强安全性
            normalized_path = os.path.normpath(os.path.abspath(directory))
            self.dest_dir.set(normalized_path)
            
    def refresh_tree(self, display_mode_override=None):
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，暂不支持刷新目录树。")  # 20260402 100551 任务运行期间禁止刷新目录树以避免状态混乱
            return
        # 清空树形视图
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.checked_items.clear()
        self.ui_update_counter = 0
        self.tree_item_paths = {}  # 20260402 085901 刷新时重置路径映射
        self.lazy_loaded_items = set()  # 20260402 085901 刷新时重置懒加载状态
        self.current_tree_display_mode_override = display_mode_override  # 20260402 085901 记录显示模式覆盖值
            
        # 添加根目录
        source_path = self.source_dir.get()
        if source_path:
            display_mode = display_mode_override or self.tree_display_mode.get()
            if self.show_directory_sizes.get() and display_mode != "files_only":
                self._build_directory_size_cache(source_path, display_mode)  # 20260402 085901 仅在需要显示目录大小时构建缓存
            else:
                self.directory_size_cache = {}  # 20260402 085901 关闭目录大小显示时清空缓存避免误用
            self.add_directory_to_tree('', source_path, display_mode_override)  # 20260402 085901 懒加载：仅加载当前层级
            self._process_pending_ui(force=True)
            
    def add_directory_to_tree(self, parent, path, display_mode_override=None):
        try:
            # 检查路径是否可访问
            if not os.access(path, os.R_OK):
                # 如果无法访问，添加一个提示节点
                self.tree.insert(
                    parent, 'end', text=f"[无法访问: {os.path.basename(path)}]",
                    values=("", "受限目录", "N/A")
                )
                return

            mode = self.copy_mode.get()
            display_mode = display_mode_override or self.tree_display_mode.get()
            with os.scandir(path) as iterator:  # 20260402 085901 目录树懒加载：仅加载当前层级
                for entry in iterator:
                    try:
                        item = entry.name
                        full_path = entry.path
                        is_dir = entry.is_dir(follow_symlinks=False)
                        item_type = "目录" if is_dir else "文件"

                        if display_mode == "files_only":
                            if is_dir:
                                continue
                        elif display_mode == "directories_only":
                            if not is_dir:
                                continue
                        elif mode not in ["custom", "export_names"] and not is_dir:
                            continue

                        show_checkbox = mode in ["custom", "selected_levels"]
                        if mode == "selected_levels" and not is_dir:
                            show_checkbox = False

                        size_str = self.get_size_in_kb(full_path)

                        item_id = self.tree.insert(
                            parent, 'end', text=item,
                            values=("🔲" if show_checkbox else "", item_type, size_str)
                        )
                        self.tree_item_paths[item_id] = full_path  # 20260402 085901 记录节点路径映射

                        row_count = len(self.tree.get_children(parent))
                        tag = 'oddrow' if row_count % 2 == 1 else 'evenrow'
                        self.tree.item(item_id, tags=(tag,))

                        if is_dir and display_mode != "files_only":
                            self.tree.insert(item_id, 'end', text=self.lazy_placeholder_text, values=("", "", ""))  # 20260402 085901 插入占位子节点以支持展开

                        self._process_pending_ui()
                    except (PermissionError, OSError):
                        continue
                    
        except (PermissionError, OSError) as e:
            # 整个目录访问失败
            if parent == '':
                # 如果是根目录访问失败，显示错误消息
                messagebox.showerror("错误", f"无法访问选定的目录:\n{path}\n\n错误信息: {str(e)}\n\n请选择一个有访问权限的目录。")
            else:
                # 子目录访问失败，添加提示节点
                self.tree.insert(
                    parent, 'end', text=f"[无法访问: {os.path.basename(path)}]",
                    values=("", "受限目录", "N/A")
                )
        except Exception as e:
            # 其他未预期的错误
            messagebox.showerror("错误", f"访问目录时出现未知错误: {str(e)}")

    def on_tree_open(self, event):
        item_id = self.tree.focus()  # 20260402 085901 展开节点时触发懒加载
        if not item_id:
            return
        if item_id in self.lazy_loaded_items:
            return
        children = self.tree.get_children(item_id)
        if len(children) == 1 and self.tree.item(children[0]).get('text') == self.lazy_placeholder_text:
            self.tree.delete(children[0])  # 20260402 085901 删除占位子节点
        else:
            if children:
                self.lazy_loaded_items.add(item_id)  # 20260402 085901 节点已加载过（或有真实子节点）
                return
        dir_path = self.tree_item_paths.get(item_id)
        if not dir_path:
            return
        try:
            self.add_directory_to_tree(item_id, dir_path, self.current_tree_display_mode_override)  # 20260402 085901 仅在展开时加载下一层
            self.lazy_loaded_items.add(item_id)  # 20260402 085901 标记节点已完成加载
        except Exception:
            self.tree.insert(item_id, 'end', text=f"[无法访问: {os.path.basename(dir_path)}]", values=("", "受限目录", "N/A"))  # 20260402 085901 展开失败时显示提示节点

    # 20260906 205815 新增：目录树搜索功能
    def _search_in_tree(self):
        """按文件名/目录名搜索源目录，弹出结果对话框"""
        keyword = self.search_var.get().strip()
        if not keyword:
            self._show_warning_message("提示", "请输入搜索关键词！")
            return
        if not self._require_source_directory():
            return

        source_path = self.source_dir.get()
        keyword_lower = keyword.lower()
        results = []
        try:
            for dirpath, dirnames, filenames in os.walk(source_path):
                for name in dirnames:
                    if keyword_lower in name.lower():
                        results.append({"name": name, "type": "目录", "path": os.path.join(dirpath, name)})
                for name in filenames:
                    if keyword_lower in name.lower():
                        results.append({"name": name, "type": "文件", "path": os.path.join(dirpath, name)})
                self._process_pending_ui()
        except Exception as e:
            self._show_error_message("错误", f"搜索过程中出错:\n{str(e)}")
            return

        if not results:
            self.search_result_var.set(f"未找到匹配项")
            self._show_info_message("搜索结果", f"未找到包含 \"{keyword}\" 的文件或目录。")
            return

        self.search_result_var.set(f"找到 {len(results)} 个匹配项")
        self._show_search_results(results, keyword)

    def _show_search_results(self, results, keyword):
        """显示搜索结果对话框"""
        result_dialog = tk.Toplevel(self.root)
        result_dialog.title(f"搜索结果 - \"{keyword}\"")
        result_dialog.transient(self.root)
        result_dialog.grab_set()

        main_frame = ttk.Frame(result_dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main_frame, text=f"共找到 {len(results)} 个匹配项（关键词：{keyword}）",
                  font=("TkDefaultFont", 11, "bold")).pack(pady=(0, 10))

        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        result_tree = ttk.Treeview(tree_frame, columns=("name", "type", "path"), show="headings")
        result_tree.heading("name", text="名称")
        result_tree.heading("type", text="类型")
        result_tree.heading("path", text="路径")
        result_tree.column("name", width=200)
        result_tree.column("type", width=60, anchor='center')
        result_tree.column("path", width=400)

        v_sb = ttk.Scrollbar(tree_frame, orient="vertical", command=result_tree.yview)
        result_tree.configure(yscrollcommand=v_sb.set)
        result_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_sb.pack(side=tk.RIGHT, fill=tk.Y)

        for item in results:
            result_tree.insert("", "end", values=(item["name"], item["type"], item["path"]))

        result_tree.bind("<Double-1>", lambda e: self._locate_search_result(result_tree, result_dialog))

        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(button_frame, text="在树中定位",
                   command=lambda: self._locate_search_result(result_tree, result_dialog)).pack(side=tk.LEFT)
        ttk.Button(button_frame, text="复制结果",
                   command=lambda: self._copy_search_results(result_tree)).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(button_frame, text="关闭", command=result_dialog.destroy).pack(side=tk.RIGHT)

        result_dialog.update_idletasks()
        req_width = main_frame.winfo_reqwidth() + 30
        req_height = main_frame.winfo_reqheight() + 30
        dw = max(680, min(req_width, result_dialog.winfo_screenwidth() - 100))
        dh = max(350, min(req_height, int(result_dialog.winfo_screenheight() * 0.8)))
        x = (result_dialog.winfo_screenwidth() - dw) // 2
        y = max(50, (result_dialog.winfo_screenheight() - dh) // 3)
        result_dialog.geometry(f"{dw}x{dh}+{x}+{y}")

    def _locate_search_result(self, result_tree, result_dialog):
        """将选中的搜索结果在目录树中定位（展开并选中）"""
        sel = result_tree.selection()
        if not sel:
            self._show_warning_message("提示", "请先选择一条搜索结果！")
            return
        values = result_tree.item(sel[0], "values")
        target_path = values[2]
        if self._expand_to_path(target_path):
            result_dialog.destroy()

    def _expand_to_path(self, target_path):
        """在目录树中展开到指定路径并选中（自动触发懒加载）"""
        source_path = self.source_dir.get()
        try:
            rel = os.path.relpath(target_path, source_path)
        except ValueError:
            self._show_warning_message("提示", "目标路径不在源目录下，无法定位。")
            return False

        parts = [p for p in rel.split(os.sep) if p]
        if not parts:
            return False

        parent = ''
        current_path = os.path.normpath(source_path)
        for part in parts:
            found = None
            target_norm = os.path.normcase(os.path.normpath(os.path.join(current_path, part)))
            for child in self.tree.get_children(parent):
                child_path = self.tree_item_paths.get(child)
                if child_path and os.path.normcase(os.path.normpath(child_path)) == target_norm:
                    found = child
                    break
            if found is None:
                self._show_warning_message("提示", f"在目录树中未找到：{part}\n请先刷新目录树后再搜索。")
                return False
            # 如果是目录且未加载，手动展开（触发懒加载）
            if found not in self.lazy_loaded_items:
                children = self.tree.get_children(found)
                if len(children) == 1 and self.tree.item(children[0]).get('text') == self.lazy_placeholder_text:
                    self.tree.delete(children[0])
                    dir_path = self.tree_item_paths.get(found)
                    if dir_path:
                        try:
                            self.add_directory_to_tree(found, dir_path, self.current_tree_display_mode_override)
                            self.lazy_loaded_items.add(found)
                        except Exception:
                            pass
                else:
                    self.lazy_loaded_items.add(found)
            self.tree.item(found, open=True)
            parent = found
            current_path = os.path.normpath(os.path.join(current_path, part))

        self.tree.see(parent)
        self.tree.focus(parent)
        self.tree.selection_set(parent)
        return True

    def _copy_search_results(self, result_tree):
        """复制搜索结果到剪贴板"""
        lines = ["名称\t类型\t路径"]
        for item_id in result_tree.get_children(""):
            values = result_tree.item(item_id, "values")
            lines.append(f"{values[0]}\t{values[1]}\t{values[2]}")
        text = "\n".join(lines)
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.root.update()
        self._show_info_message("完成", f"已复制 {len(lines) - 1} 条搜索结果到剪贴板。")

    def start_copy(self):
        if not self.validate_inputs():
            return
            
        try:
            mode = self.copy_mode.get()
            if mode == "single_level":
                self.copy_single_level()
            elif mode == "selected_levels":
                self.copy_selected_levels()
            elif mode == "all_levels":
                self.copy_all_levels()
            elif mode == "export_names":
                self.export_names()
            else:
                self.copy_custom()

            if getattr(self, "task_running", False):
                return  # 20260402 093200 分片任务执行中，由任务回调提示结果，避免提前弹出“完成”
                
            if mode == "export_names":
                messagebox.showinfo("完成", "导出操作完成!")
                pass  # 状态标签已删除
            else:
                messagebox.showinfo("完成", "复制操作完成!")
                pass  # 状态标签已删除
        except Exception as e:
            if self.copy_mode.get() == "export_names":
                messagebox.showerror("错误", f"导出过程中出错:\n{str(e)}")
                pass  # 状态标签已删除
            else:
                messagebox.showerror("错误", f"复制过程中出错:\n{str(e)}")
                pass  # 状态标签已删除
        finally:
            pass
            
    def validate_inputs(self):
        if not self.source_dir.get() or not self.dest_dir.get():
            messagebox.showwarning("警告", "请选择源目录和目标目录!")
            return False
        if self.source_dir.get() == self.dest_dir.get():
            messagebox.showwarning("警告", "源目录和目标目录不能相同!")
            return False
        return True
        
    def copy_single_level(self):
        """仅复制一层目录（空目录）"""
        src = self.source_dir.get()
        dst = self.dest_dir.get()
        try:
            dirs = []
            with os.scandir(src) as it:  # 20260402 091600 使用scandir构建队列
                for entry in it:
                    if entry.is_dir(follow_symlinks=False):
                        dirs.append(entry.name)
        except Exception as e:
            messagebox.showerror("错误", f"无法读取目录:\n{src}\n{str(e)}")
            return
        
        if not dirs:
            self._show_info_message("提示", "源目录下没有子目录，无需复制。")  # 20260828 111834 无子目录时给出明确提示
            return
        
        # 分片执行 + 可取消
        self.task_runner = self.TaskRunner(self, description="复制一层目录", total=len(dirs))  # 20260402 091600 启动任务
        
        def step_fn(dir_name, batch_ops):
            target_path = os.path.join(dst, dir_name)
            if os.path.exists(target_path):
                return "skipped"
            os.makedirs(target_path, exist_ok=True)
            batch_ops.append({  # 20260402 091600 记录子操作
                'type': 'copy',
                'details': {
                    'source_path': os.path.join(src, dir_name),
                    'target_path': target_path,
                    'operation_type': 'copy_single_level',
                    'is_directory': True
                }
            })
        
        def on_done(summary):
            if summary.get("batch_operations"):
                desc = "复制一层目录" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {  # 20260402 091600 写入批次历史
                    'description': desc,
                    'operation_type': 'copy_single_level',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()
            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {summary.get('completed',0)} 项，跳过 {summary.get('skipped',0)} 项，失败 {summary.get('failed',0)} 项。")  # 20260402 091600 取消提示
            else:
                self._show_info_message("完成", "复制操作完成!")  # 20260402 093200 分片任务完成后统一提示
        
        self.task_runner.start(dirs, step_fn, on_done)  # 20260402 091600 启动分片任务
            
    def collect_structure_with_path(self, path, structure_data, level, root_path):
        """递归收集目录结构信息（包含完整路径）"""
        try:
            # 检查路径是否可访问
            if not os.access(path, os.R_OK):
                structure_data.append({
                    'name': f"[无法访问: {os.path.basename(path)}]",
                    'level': level,
                    'is_dir': False,
                    'full_path': path
                })
                return
                
            items = sorted(os.listdir(path))
            for item in items:
                full_path = os.path.join(path, item)
                
                try:
                    # 检查单个项目是否可访问
                    is_dir = os.path.isdir(full_path)
                    structure_data.append({
                        'name': item,
                        'level': level,
                        'is_dir': is_dir,
                        'full_path': full_path
                    })
                    
                    if is_dir:
                        self.collect_structure_with_path(full_path, structure_data, level + 1, root_path)
                    self._process_pending_ui()
                        
                except (PermissionError, OSError):
                    # 单个文件/目录访问失败，添加提示信息但继续处理
                    structure_data.append({
                        'name': f"[无法访问: {item}]",
                        'level': level,
                        'is_dir': False,
                        'full_path': full_path
                    })
                    
        except (PermissionError, OSError):
            # 整个目录访问失败
            structure_data.append({
                'name': f"[无法访问: {os.path.basename(path)}]",
                'level': level,
                'is_dir': False,
                'full_path': path
            })
        except Exception as e:
            # 其他未预期的错误
            structure_data.append({
                'name': f"[错误: {os.path.basename(path)} - {str(e)}]",
                'level': level,
                'is_dir': False,
                'full_path': path
            })
            
            self._process_pending_ui()
            
    def copy_selected_levels(self):
        """复制用户勾选的目录（仅结构）"""
        if not self.checked_items:
            messagebox.showwarning("警告", "请选择要复制的目录!")
            return
            
        self._reset_runtime_warnings()
        source_root = self.source_dir.get()  # 20260402 100551 固化源/目标与路径队列，避免任务期间树刷新导致路径解析异常
        dest_root = self.dest_dir.get()  # 20260402 100551 固化源/目标与路径队列，避免任务期间树刷新导致路径解析异常
        items = []
        for item_id in list(self.checked_items):
            try:
                item_path = self.get_item_path(item_id)
            except Exception:
                self._add_runtime_warning("选中项路径解析失败，已跳过。")  # 20260402 100551 防止树刷新/节点失效导致异常
                continue
            if not os.path.isdir(item_path):
                continue
            try:
                rel_path = os.path.relpath(item_path, source_root)
            except Exception:
                continue
            dst_path = os.path.join(dest_root, rel_path)
            items.append((item_path, dst_path, rel_path))

        if not items:
            return

        self.task_runner = self.TaskRunner(self, description="复制选定层级目录", total=len(items))  # 20260402 093200 启动任务

        def step_fn(item, batch_ops):
            item_path, dst_path, rel_path = item
            if os.path.exists(dst_path):
                return "skipped"

            try:
                os.makedirs(dst_path, exist_ok=True)
                batch_ops.append({  # 20260402 093200 记录子操作
                    'type': 'copy',
                    'details': {
                        'source_path': item_path,
                        'target_path': dst_path,
                        'operation_type': 'copy_selected_levels',
                        'is_directory': True
                    }
                })
            except Exception:
                self._add_runtime_warning(f"创建目录失败：{rel_path}")  # 20260402 093200 记录失败原因用于汇总提示
                raise

        def on_done(summary):
            if summary.get("batch_operations"):
                desc = "复制选定层级目录" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {  # 20260402 093200 写入批次历史
                    'description': desc,
                    'operation_type': 'copy_selected_levels',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()
            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {summary.get('completed',0)} 项，跳过 {summary.get('skipped',0)} 项，失败 {summary.get('failed',0)} 项。")  # 20260402 093200 取消提示
            else:
                self._show_info_message("完成", "复制操作完成!")  # 20260402 093200 分片任务完成后统一提示
            self._show_runtime_warnings("复制提示")  # 20260402 093200 汇总提示

        self.task_runner.start(items, step_fn, on_done)  # 20260402 093200 启动分片任务
                
    def copy_all_levels(self):
        """复制完整目录树（空目录）"""
        self._reset_runtime_warnings()
        src_root = self.source_dir.get()
        dst_root = self.dest_dir.get()
        failed_prefixes = []  # 20260402 093522 记录创建失败的目标路径前缀，避免对子目录反复报错
        self.task_runner = self.TaskRunner(self, description="复制所有层级目录")  # 20260402 100551 将扫描纳入分片执行，避免预扫描卡顿
        
        def step_fn(item, batch_ops):
            src_dir, dst_dir = item
            new_items = []
            dst_norm = os.path.normcase(os.path.abspath(dst_dir))
            for prefix in failed_prefixes:
                if dst_norm.startswith(prefix):
                    return ("skipped", new_items)

            created = False
            try:
                if not os.path.exists(dst_dir):
                    os.makedirs(dst_dir, exist_ok=True)
                    created = True
            except Exception:
                failed_prefixes.append(dst_norm + os.sep)
                self._add_runtime_warning(f"无法创建目录：{dst_dir}")
                return ("failed", new_items)  # 20260402 100551 创建失败不再抛异常，避免整体中断

            if created:
                batch_ops.append({  # 20260402 093522 记录子操作（目录创建）
                    'type': 'copy',
                    'details': {
                        'source_path': src_dir,
                        'target_path': dst_dir,
                        'operation_type': 'copy_all_levels',
                        'is_directory': True
                    }
                })

            try:
                with os.scandir(src_dir) as iterator:  # 20260402 100551 分片扫描子目录
                    for entry in iterator:
                        if entry.is_dir(follow_symlinks=False):
                            new_items.append((entry.path, os.path.join(dst_dir, entry.name)))
            except Exception:
                self._add_runtime_warning(f"无法访问目录：{src_dir}")
                return ("skipped", [])

            return (("processed" if created else "skipped"), new_items)  # 20260402 100551 未创建但继续扫描子目录时计为跳过
        
        def on_done(summary):
            if summary.get("batch_operations"):
                desc = "复制所有层级目录" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {  # 20260402 093522 写入批次历史
                    'description': desc,
                    'operation_type': 'copy_all_levels',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()
            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {summary.get('completed',0)} 项，跳过 {summary.get('skipped',0)} 项，失败 {summary.get('failed',0)} 项。")  # 20260402 093522 取消提示
            else:
                self._show_info_message("完成", "复制操作完成!")  # 20260402 093522 分片任务完成后统一提示
            self._show_runtime_warnings("复制提示")  # 20260402 093522 汇总提示
        
        self.task_runner.start([(src_root, dst_root)], step_fn, on_done)  # 20260402 100551 启动分片任务（从根目录开始）
        
    def copy_dir_tree(self, src, dst, batch_operations=None):
        try:
            if not os.path.exists(dst):
                os.makedirs(dst, exist_ok=True)
                if batch_operations is not None:
                    batch_operations.append({  # 20260402 085047 记录子操作（目录创建）
                        'type': 'copy',
                        'details': {
                            'source_path': src,
                            'target_path': dst,
                            'operation_type': 'copy_all_levels',
                            'is_directory': True
                        }
                    })
                    
        except (PermissionError, OSError) as e:
            self._add_runtime_warning(f"无法创建目录：{dst}")
            return
        
        try:
            items = os.listdir(src)
        except (PermissionError, OSError) as e:
            self._add_runtime_warning(f"无法访问目录：{src}")
            return
        
        for item in items:
            src_item = os.path.join(src, item)
            try:
                if os.path.isdir(src_item):
                    dst_item = os.path.join(dst, item)
                    self.copy_dir_tree(src_item, dst_item, batch_operations)
            except (PermissionError, OSError) as e:
                self._add_runtime_warning(f"跳过无法访问的目录：{src_item}")
                continue
            
            self._process_pending_ui()
                
    def copy_custom(self):
        """自定义复制（允许选择文件和目录）"""
        if not self.checked_items:
            messagebox.showwarning("警告", "请选择要复制的项目!")
            return
            
        self._reset_runtime_warnings()
        source_root = self.source_dir.get()  # 20260402 100551 固化源/目标与路径队列，避免任务期间树刷新导致路径解析异常
        dest_root = self.dest_dir.get()  # 20260402 100551 固化源/目标与路径队列，避免任务期间树刷新导致路径解析异常
        items = []
        for item_id in list(self.checked_items):
            try:
                item_path = self.get_item_path(item_id)
            except Exception:
                self._add_runtime_warning("选中项路径解析失败，已跳过。")  # 20260402 100551 防止树刷新/节点失效导致异常
                continue
            if not os.path.exists(item_path):
                continue
            try:
                rel_path = os.path.relpath(item_path, source_root)
            except Exception:
                continue
            dst_path = os.path.join(dest_root, rel_path)
            items.append((item_path, dst_path, rel_path, os.path.isdir(item_path)))

        if not items:
            return

        self.task_runner = self.TaskRunner(self, description="自定义复制", total=len(items))  # 20260402 094200 启动任务
        
        def step_fn(item, batch_ops):
            item_path, dst_path, rel_path, is_dir = item
            try:
                if is_dir:
                    if os.path.exists(dst_path):
                        return "skipped"
                    shutil.copytree(item_path, dst_path, dirs_exist_ok=True)
                    batch_ops.append({  # 20260402 094200 记录子操作
                        'type': 'copy',
                        'details': {
                            'source_path': item_path,
                            'target_path': dst_path,
                            'operation_type': 'copy_custom',
                            'is_directory': True
                        }
                    })
                else:
                    if os.path.exists(dst_path):
                        return "skipped"
                    os.makedirs(os.path.dirname(dst_path), exist_ok=True)
                    shutil.copy2(item_path, dst_path)
                    batch_ops.append({  # 20260402 094200 记录子操作
                        'type': 'copy',
                        'details': {
                            'source_path': item_path,
                            'target_path': dst_path,
                            'operation_type': 'copy_custom',
                            'is_directory': False
                        }
                    })
            except Exception as e:
                # 20260828 111834 批量复制失败不再逐项弹窗，改为收集后由任务结束统一汇总提示
                self._add_runtime_warning(f"复制失败：{rel_path}（{str(e)}）")
                raise
        
        def on_done(summary):
            if summary.get("batch_operations"):
                desc = "自定义复制" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {  # 20260402 094200 写入批次历史
                    'description': desc,
                    'operation_type': 'copy_custom',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()
            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {summary.get('completed',0)} 项，跳过 {summary.get('skipped',0)} 项，失败 {summary.get('failed',0)} 项。")  # 20260402 094200 取消提示
            else:
                self._show_info_message("完成", "复制操作完成!")  # 20260402 094200 分片任务完成后统一提示
            self._show_runtime_warnings("复制提示")  # 20260402 094200 汇总提示

        self.task_runner.start(items, step_fn, on_done)  # 20260402 094200 启动分片任务
                
    def export_names(self):
        """导出目录和文件名称到指定格式的文档"""
        src = self.source_dir.get()
        dst = self.dest_dir.get()
        export_format = self.export_format.get()
        self._save_export_format()  # 20260906 205815 记忆用户选择的导出格式
        
        # 生成导出文件名（格式：导出时间+目录名+'目录结构'）
        from datetime import datetime
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        source_name = os.path.basename(src)
        
        # 根据格式确定文件扩展名
        if export_format in ["html", "html_link"]:
            file_ext = "html"
        elif export_format in ["md", "md_link"]:
            file_ext = "md"
        elif export_format in ["docx", "docx_link"]:
            file_ext = "docx"
        elif export_format in ["xlsx", "xlsx_link"]:
            file_ext = "xlsx"  # 20260402 103300 修复xlsx_link导出扩展名，统一为.xlsx
        elif export_format == "json":
            file_ext = "json"  # 20260402 103300 明确json扩展名
        else:
            file_ext = export_format

        filename_prefix = f"{current_time}_{source_name}_目录结构"
        if export_format == "xlsx_link":
            filename_prefix += "_含链接"  # 20260402 103300 避免xlsx与xlsx(含链接)导出同名覆盖
        filename = f"{filename_prefix}.{file_ext}"
        output_path = os.path.join(dst, filename)
        
        # 20260828 111834 导出改为“分片收集目录结构 + 生成文档”流程，避免大目录同步遍历卡界面，并支持“停止当前任务”取消
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，请稍后再试。")
            return

        structure_data = []
        self._reset_runtime_warnings()
        self.task_running = True
        self.cancel_requested = False
        self._export_collecting = True  # 20260828 111834 标记导出收集阶段，供“停止当前任务”取消
        self._set_task_ui_state(True)
        self.cancel_task_button.configure(state="normal")  # 20260828 111834 启用“停止当前任务”以支持取消导出
        self.progress_var.set("正在收集目录结构...")

        def on_collect_done(cancelled):
            self._export_collecting = False
            self.task_running = False
            self._set_task_ui_state(False)
            self.cancel_task_button.configure(state="disabled")  # 20260828 111834 收集结束恢复按钮状态
            if cancelled:
                self.progress_var.set("")
                self._show_info_message("提示", "导出已取消。")
                return
            self.progress_var.set("正在生成导出文档...")
            try:
                self._finish_export(structure_data, source_name, export_format, output_path)
                self._show_info_message("完成", "导出操作完成!")
            except Exception as e:
                self._show_error_message("错误", f"导出过程中出错:\n{str(e)}")
            finally:
                self.progress_var.set("")

        self._start_export_collect(src, structure_data, on_collect_done)

    def _start_export_collect(self, root_path, structure_data, on_done):
        """分片收集目录结构（显式栈深度优先，顺序与原递归逻辑一致），期间界面可响应、可取消"""
        stack = [(root_path, 0)]
        self._export_after_id = None

        def cleanup_after():
            if self._export_after_id is not None:
                try:
                    self.root.after_cancel(self._export_after_id)
                except Exception:
                    pass
                self._export_after_id = None

        def process_slice():
            if self.cancel_requested:
                cleanup_after()
                on_done(True)
                return
            count = 0
            while stack and count < self.ui_update_interval:
                path, level = stack.pop()
                count += 1
                try:
                    if not os.access(path, os.R_OK):
                        structure_data.append({'name': f"[无法访问: {os.path.basename(path)}]", 'level': level, 'is_dir': False, 'full_path': path})
                        continue
                    items = sorted(os.listdir(path))
                except (PermissionError, OSError):
                    structure_data.append({'name': f"[无法访问: {os.path.basename(path)}]", 'level': level, 'is_dir': False, 'full_path': path})
                    continue
                except Exception as e:
                    structure_data.append({'name': f"[错误: {os.path.basename(path)} - {str(e)}]", 'level': level, 'is_dir': False, 'full_path': path})
                    continue
                subdirs = []
                for item in items:
                    full_path = os.path.join(path, item)
                    try:
                        is_dir = os.path.isdir(full_path)
                    except (PermissionError, OSError):
                        is_dir = False
                    structure_data.append({'name': item, 'level': level, 'is_dir': is_dir, 'full_path': full_path})
                    if is_dir:
                        subdirs.append(full_path)
                # 倒序压栈以保持与原递归一致的深度优先顺序
                for sub in reversed(subdirs):
                    stack.append((sub, level + 1))
            if stack:
                self.progress_var.set(f"正在收集目录结构...（剩余 {len(stack)} 个目录）")
                self._export_after_id = self.root.after(0, process_slice)
            else:
                cleanup_after()
                on_done(False)

        process_slice()

    def _finish_export(self, structure_data, source_name, export_format, output_path):
        """根据收集到的结构数据生成并写入导出文件（原 export_names 后半部分逻辑）"""
        if export_format == "txt":
            content = self.generate_txt_content(structure_data, source_name)
        elif export_format == "html":
            content = self.generate_html_content(structure_data, source_name)
        elif export_format == "html_link":
            content = self.generate_html_content_with_links(structure_data, source_name)
        elif export_format == "md":
            content = self.generate_md_content(structure_data, source_name)
        elif export_format == "md_link":
            content = self.generate_md_content_with_links(structure_data, source_name)
        elif export_format == "docx":
            self.generate_docx_content(structure_data, source_name, output_path)
            return  # DOCX直接写入文件，不需要后续的文本写入操作
        elif export_format == "docx_link":
            self.generate_docx_content_with_links(structure_data, source_name, output_path)
            return  # DOCX直接写入文件，不需要后续的文本写入操作
        elif export_format == "json":  # 20260402 102300 新增JSON导出
            content = self.generate_json_content(structure_data, source_name)
        elif export_format == "xlsx":  # 20260402 102300 新增XLSX导出
            self.write_xlsx_file(structure_data, source_name, output_path, with_links=False)
            return  # XLSX直接写入文件
        elif export_format == "xlsx_link":  # 20260402 102300 新增XLSX(含链接)导出
            self.write_xlsx_file(structure_data, source_name, output_path, with_links=True)
            return  # XLSX直接写入文件

        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def generate_json_content(self, structure_data, source_name):  # 20260402 102300 新增：生成JSON内容
        import json
        rows = []
        for item in structure_data:
            rows.append({
                "level": int(item.get("level", 0)),
                "type": "目录" if item.get("is_dir") else "文件",
                "name": str(item.get("name", "")),
                "full_path": str(item.get("full_path", "")),
                "is_dir": bool(item.get("is_dir", False))
            })
        return json.dumps(rows, ensure_ascii=False, indent=2)
    
    def write_xlsx_file(self, structure_data, source_name, output_path, with_links=False):  # 20260402 102300 新增：写入XLSX文件（可含链接）
        try:
            import xlsxwriter
        except ImportError:
            self._show_error_message(
                "缺少依赖",
                "未检测到 xlsxwriter 库，无法导出为 .xlsx。\n\n请先安装：\n\npip install XlsxWriter"
            )
            return
        
        try:
            workbook = xlsxwriter.Workbook(output_path)
            ws = workbook.add_worksheet("目录结构")
            
            header_fmt = workbook.add_format({"bold": True})
            wrap_fmt = workbook.add_format({"text_wrap": True})
            
            ws.write(0, 0, "层级", header_fmt)
            ws.write(0, 1, "类型", header_fmt)
            ws.write(0, 2, "名称", header_fmt)
            ws.write(0, 3, "完整路径", header_fmt)
            
            row = 1
            for item in structure_data:
                level = int(item.get("level", 0))
                is_dir = bool(item.get("is_dir", False))
                name = str(item.get("name", ""))
                full_path = str(item.get("full_path", ""))
                typ = "目录" if is_dir else "文件"
                
                ws.write_number(row, 0, level)
                ws.write(row, 1, typ)
                ws.write(row, 2, name, wrap_fmt)
                
                if with_links and full_path:
                    try:
                        uri = self.path_to_file_uri(full_path)
                        ws.write_url(row, 3, uri, string=full_path)  # 链接文本用完整路径  # 20260402 102300
                    except Exception:
                        ws.write(row, 3, full_path, wrap_fmt)
                else:
                    ws.write(row, 3, full_path, wrap_fmt)
                
                row += 1
            
            ws.set_column(0, 0, 8)   # 层级
            ws.set_column(1, 1, 8)   # 类型
            ws.set_column(2, 2, 40)  # 名称
            ws.set_column(3, 3, 60)  # 完整路径
        finally:
            try:
                workbook.close()
            except Exception:
                pass
        
    def collect_structure(self, path, structure_data, level=0):
        """递归收集目录结构信息"""
        try:
            # 检查路径是否可访问
            if not os.access(path, os.R_OK):
                structure_data.append({
                    'name': f"[无法访问: {os.path.basename(path)}]",
                    'level': level,
                    'is_dir': False
                })
                return
                
            items = sorted(os.listdir(path))
            for item in items:
                full_path = os.path.join(path, item)
                
                try:
                    # 检查单个项目是否可访问
                    is_dir = os.path.isdir(full_path)
                    structure_data.append({
                        'name': item,
                        'level': level,
                        'is_dir': is_dir
                    })
                    
                    if is_dir:
                        self.collect_structure(full_path, structure_data, level + 1)
                        
                except (PermissionError, OSError):
                    # 单个文件/目录访问失败，添加提示信息但继续处理
                    structure_data.append({
                        'name': f"[无法访问: {item}]",
                        'level': level,
                        'is_dir': False
                    })
                    
        except (PermissionError, OSError):
            # 整个目录访问失败
            structure_data.append({
                'name': f"[无法访问: {os.path.basename(path)}]",
                'level': level,
                'is_dir': False
            })
        except Exception as e:
            # 其他未预期的错误
            structure_data.append({
                'name': f"[错误: {os.path.basename(path)} - {str(e)}]",
                'level': level,
                'is_dir': False
            })
            
    def generate_txt_content(self, structure_data, source_name):
        """生成TXT格式内容"""
        lines = []
        lines.append(f"目录结构导出 - {source_name}")
        lines.append("=" * 50)
        lines.append("")
        
        for item in structure_data:
            indent = "  " * item['level']
            prefix = "📁 " if item['is_dir'] else "📄 "
            lines.append(f"{indent}{prefix}{item['name']}")
            
        return "\n".join(lines)
        
    def generate_html_content(self, structure_data, source_name):
        """生成HTML格式内容"""
        lines = []
        lines.append("<!DOCTYPE html>")
        lines.append("<html lang='zh-CN'>")
        lines.append("<head>")
        lines.append("    <meta charset='UTF-8'>")
        lines.append("    <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
        lines.append(f"    <title>目录结构 - {source_name}</title>")
        lines.append("    <style>")
        lines.append("        body { font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 20px; }")
        lines.append("        h1 { color: #333; border-bottom: 2px solid #007acc; padding-bottom: 10px; }")
        lines.append("        .tree { font-family: 'Courier New', monospace; line-height: 1.6; }")
        lines.append("        .folder { color: #007acc; font-weight: bold; }")
        lines.append("        .file { color: #666; }")
        lines.append("        .indent { margin-left: 20px; }")
        lines.append("    </style>")
        lines.append("</head>")
        lines.append("<body>")
        lines.append(f"    <h1>目录结构 - {source_name}</h1>")
        lines.append("    <div class='tree'>")
        
        for item in structure_data:
            indent_class = f"indent" if item['level'] > 0 else ""
            style_class = "folder" if item['is_dir'] else "file"
            icon = "📁" if item['is_dir'] else "📄"
            
            indent_style = f"margin-left: {item['level'] * 20}px;"
            lines.append(f"        <div class='{style_class}' style='{indent_style}'>{icon} {item['name']}</div>")
            
        lines.append("    </div>")
        lines.append("</body>")
        lines.append("</html>")
        
        return "\n".join(lines)
        
    def generate_md_content(self, structure_data, source_name):
        """生成Markdown格式内容"""
        lines = []
        lines.append(f"# 目录结构 - {source_name}")
        lines.append("")
        lines.append("```")
        
        for item in structure_data:
            indent = "  " * item['level']
            prefix = "📁 " if item['is_dir'] else "📄 "
            lines.append(f"{indent}{prefix}{item['name']}")
            
        lines.append("```")
        lines.append("")
        lines.append(f"*导出时间: {self.get_current_time()}*")
        
        return "\n".join(lines)
        
    def generate_html_content_with_links(self, structure_data, source_name):
        """生成HTML格式内容（含硬链接）"""
        lines = []
        lines.append("<!DOCTYPE html>")
        lines.append("<html lang='zh-CN'>")
        lines.append("<head>")
        lines.append("    <meta charset='UTF-8'>")
        lines.append("    <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
        lines.append(f"    <title>目录结构(含链接) - {source_name}</title>")
        lines.append("    <style>")
        lines.append("        body { font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 20px; }")
        lines.append("        h1 { color: #333; border-bottom: 2px solid #007acc; padding-bottom: 10px; }")
        lines.append("        .tree { font-family: 'Courier New', monospace; line-height: 1.6; }")
        lines.append("        .folder { color: #007acc; font-weight: bold; }")
        lines.append("        .file { color: #666; }")
        lines.append("        .folder a, .file a { text-decoration: none; color: inherit; }")
        lines.append("        .folder a:hover, .file a:hover { text-decoration: underline; }")
        lines.append("        .indent { margin-left: 20px; }")
        lines.append("    </style>")
        lines.append("</head>")
        lines.append("<body>")
        lines.append(f"    <h1>目录结构(含链接) - {source_name}</h1>")
        lines.append("    <div class='tree'>")
        
        for item in structure_data:
            indent_class = f"indent" if item['level'] > 0 else ""
            style_class = "folder" if item['is_dir'] else "file"
            icon = "📁" if item['is_dir'] else "📄"
            
            # 生成文件URI（硬链接）
            file_uri = self.path_to_file_uri(item['full_path'])
            
            indent_style = f"margin-left: {item['level'] * 20}px;"
            lines.append(f"        <div class='{style_class}' style='{indent_style}'>{icon} <a href='{file_uri}'>{item['name']}</a></div>")
            
        lines.append("    </div>")
        lines.append(f"    <p><em>导出时间: {self.get_current_time()}</em></p>")
        lines.append("</body>")
        lines.append("</html>")
        
        return "\n".join(lines)
        
    def generate_md_content_with_links(self, structure_data, source_name):
        """生成Markdown格式内容（含硬链接）"""
        lines = []
        lines.append(f"# 目录结构(含链接) - {source_name}")
        lines.append("")
        
        for item in structure_data:
            indent = "  " * item['level']
            prefix = "📁 " if item['is_dir'] else "📄 "
            
            # 生成文件URI（硬链接）
            file_uri = self.path_to_file_uri(item['full_path'])
            
            # Markdown链接格式
            lines.append(f"{indent}- {prefix}[{item['name']}]({file_uri})")
            
        lines.append("")
        lines.append(f"*导出时间: {self.get_current_time()}*")
        
        return "\n".join(lines)
        
    def path_to_file_uri(self, file_path):
        """将文件路径转换为file:// URI格式（硬链接）"""
        import urllib.parse
        
        # 规范化路径
        normalized_path = os.path.abspath(file_path)
        
        # 在Windows系统中，需要特殊处理路径格式
        if os.name == 'nt':  # Windows
            # 将反斜杠转换为正斜杠
            normalized_path = normalized_path.replace('\\', '/')
            # 确保以斜杠开头
            if not normalized_path.startswith('/'):
                normalized_path = '/' + normalized_path
        
        # URL编码路径中的特殊字符（保持中文字符可读性）
        encoded_path = urllib.parse.quote(normalized_path, safe='/:@!$&\'()*+,;=')
        
        return f"file://{encoded_path}"
        
    def generate_docx_content(self, structure_data, source_name, output_path):
        """生成DOCX格式内容"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建新文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(f'目录结构 - {source_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加导出时间
            time_para = doc.add_paragraph(f'导出时间: {self.get_current_time()}')
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 添加空行
            doc.add_paragraph()
            
            # 添加目录结构内容
            for item in structure_data:
                # 创建段落
                para = doc.add_paragraph()
                
                # 添加缩进
                para.paragraph_format.left_indent = Inches(item['level'] * 0.3)
                
                # 添加图标和文件名
                icon = "📁 " if item['is_dir'] else "📄 "
                para.add_run(f"{icon}{item['name']}")
                
            # 保存文档
            doc.save(output_path)
            
            # 显示成功消息
            messagebox.showinfo("导出成功", f"目录结构已导出到:\n{output_path}")
            
        except ImportError:
            messagebox.showerror("错误", "缺少python-docx库，请先安装:\npip install python-docx")
        except Exception as e:
            messagebox.showerror("导出失败", f"导出DOCX文件时发生错误:\n{str(e)}")
            
    def generate_docx_content_with_links(self, structure_data, source_name, output_path):
        """生成DOCX格式内容（含硬链接）"""
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # 创建新文档
            doc = Document()
            
            # 添加标题
            title = doc.add_heading(f'目录结构(含链接) - {source_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加导出时间
            time_para = doc.add_paragraph(f'导出时间: {self.get_current_time()}')
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 添加空行
            doc.add_paragraph()
            
            # 添加目录结构内容
            for item in structure_data:
                # 创建段落
                para = doc.add_paragraph()
                
                # 添加缩进
                para.paragraph_format.left_indent = Inches(item['level'] * 0.3)
                
                # 添加图标
                icon = "📁 " if item['is_dir'] else "📄 "
                para.add_run(icon)
                
                # 生成文件URI（硬链接）
                file_uri = self.path_to_file_uri(item['full_path'])
                
                # 添加超链接
                try:
                    # 使用python-docx的内置超链接功能
                    hyperlink_run = para.add_run(item['name'])
                    
                    # 创建超链接关系
                    part = para.part
                    r_id = part.relate_to(file_uri, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                    
                    # 创建超链接XML元素
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    
                    # 创建运行元素
                    new_run = OxmlElement('w:r')
                    
                    # 设置运行属性（蓝色、下划线）
                    rPr = OxmlElement('w:rPr')
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), '0000FF')
                    rPr.append(color)
                    
                    u = OxmlElement('w:u')
                    u.set(qn('w:val'), 'single')
                    rPr.append(u)
                    
                    new_run.append(rPr)
                    
                    # 添加文本
                    text_elem = OxmlElement('w:t')
                    text_elem.text = item['name']
                    new_run.append(text_elem)
                    
                    hyperlink.append(new_run)
                    
                    # 移除之前添加的普通文本运行
                    para._element.remove(hyperlink_run._element)
                    
                    # 添加超链接到段落
                    para._element.append(hyperlink)
                    
                except Exception as link_error:
                    # 如果超链接创建失败，回退到普通文本
                    hyperlink_run = para.add_run(item['name'])
                    hyperlink_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                    hyperlink_run.font.underline = True
                
            # 添加说明
            doc.add_paragraph()
            note_para = doc.add_paragraph('说明：点击文件名可以直接打开对应的文件或文件夹。')
            note_para.italic = True
            
            # 保存文档
            doc.save(output_path)
            
            # 显示成功消息
            messagebox.showinfo("导出成功", f"目录结构(含链接)已导出到:\n{output_path}")
            
        except ImportError:
            messagebox.showerror("错误", "缺少python-docx库，请先安装:\npip install python-docx")
        except Exception as e:
            messagebox.showerror("导出失败", f"导出DOCX文件时发生错误:\n{str(e)}")
            
    def get_current_time(self):
        """获取当前时间字符串"""
        import datetime
        return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
    def get_item_path(self, item_id):
        mapped_path = self.tree_item_paths.get(item_id)  # 20260402 085901 优先使用节点路径映射，提升懒加载场景路径解析可靠性
        if mapped_path:
            return mapped_path
        path_parts = []
        while item_id:
            path_parts.insert(0, self.tree.item(item_id)['text'])
            item_id = self.tree.parent(item_id)
        return os.path.join(self.source_dir.get(), *path_parts)
        
    def cancel_copy(self):
        if messagebox.askyesno("确认", "确定要清空当前勾选项吗？\n（不停止正在执行的任务；停止任务请点“停止当前任务”）"):  # 20260402 102000 调整文案并明确用途
            self.deselect_all()  # 20260402 102000 清空勾选项
            self._reset_runtime_warnings()  # 20260402 102000 同步清空运行提示

    def rename_current_level(self):
        """重命名本级文件夹"""
        if not self._require_source_directory():
            return
            
        dialog, find_var, replace_var = self.create_rename_dialog("本级目录字符替换")
        
        def do_rename():
            find_text = find_var.get()
            replace_text = replace_var.get()
            
            if not find_text:
                self._show_warning_message("警告", "请输入要查找的字符串!")
                return
            
            dialog.destroy()
            self._start_rename_current_level_replace(find_text, replace_text)  # 20260402 094929 分片执行：启动本级目录字符替换
                
        ttk.Button(dialog, text="确定", command=do_rename).pack(pady=10)

    def _start_rename_current_level_replace(self, find_text, replace_text):  # 20260402 094929 新增：本级目录字符替换（分片任务）
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，请稍后再试。")
            return

        source_path = self.source_dir.get()
        self._reset_runtime_warnings()
        candidates = []
        try:
            with os.scandir(source_path) as iterator:
                for entry in iterator:
                    if not entry.is_dir(follow_symlinks=False):
                        continue
                    name = entry.name
                    if find_text not in name:
                        continue
                    new_name = name.replace(find_text, replace_text)
                    if new_name == name:
                        continue
                    candidates.append((name, new_name))
        except Exception as e:
            self._show_error_message("错误", f"无法读取目录:\n{source_path}\n{str(e)}")
            return

        if not candidates:
            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("文件夹", 0, 0, 0, 0),
                refresh_mode="directories_only",
                warning_title="重命名提示",
                update_history=False
            )
            return

        self.task_runner = self.TaskRunner(self, description="本级目录字符替换", total=len(candidates))  # 20260402 094929 启动任务

        def step_fn(item, batch_ops):
            old_name, new_name = item
            old_path = os.path.join(source_path, old_name)
            new_path, warning_message = self._prepare_rename_target(source_path, old_name, new_name, "目录")
            if warning_message:
                self._add_runtime_warning(warning_message)
                return "skipped"
            if new_path is None:
                return "skipped"
            try:
                os.rename(old_path, new_path)
                batch_ops.append({
                    'type': 'rename',
                    'details': {
                        'old_path': old_path,
                        'new_path': new_path,
                        'operation_type': 'rename_current_level'
                    }
                })
            except Exception:
                self._add_runtime_warning(f"重命名失败：{old_name}")
                raise

        def on_done(summary):
            matched_count = len(candidates)
            renamed_count = summary.get("completed", 0)
            skipped_count = summary.get("skipped", 0)
            failed_count = summary.get("failed", 0)

            if summary.get("batch_operations"):
                desc = "本级目录字符替换" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {
                    'description': desc,
                    'operation_type': 'rename_current_level',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()

            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {renamed_count} 项，跳过 {skipped_count} 项，失败 {failed_count} 项。")

            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("文件夹", matched_count, renamed_count, skipped_count, failed_count),
                refresh_mode="directories_only",
                warning_title="重命名提示",
                update_history=renamed_count > 0
            )

        self.task_runner.start(candidates, step_fn, on_done)

    def rename_all_items(self):
        """重命名所有项目"""
        if not self._require_source_directory():
            return
            
        dialog, find_var, replace_var = self.create_rename_dialog("全部目录字符替换")
        
        def do_rename():
            find_text = find_var.get()
            replace_text = replace_var.get()
            
            if not find_text:
                self._show_warning_message("警告", "请输入要查找的字符串!")
                return
            
            dialog.destroy()
            self._start_rename_all_items_replace(find_text, replace_text)  # 20260402 094929 分片执行：启动全部目录字符替换
                
        ttk.Button(dialog, text="确定", command=do_rename).pack(pady=10)

    def _start_rename_all_items_replace(self, find_text, replace_text):  # 20260402 094929 新增：全部目录字符替换（分片任务）
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，请稍后再试。")
            return

        root_path = self.source_dir.get()
        self._reset_runtime_warnings()
        matched_count = {"value": 0}
        renamed_count = {"value": 0}
        skipped_count = {"value": 0}
        failed_count = {"value": 0}

        self.task_runner = self.TaskRunner(self, description="全部目录字符替换")  # 20260402 094929 启动任务（总量不预估）

        def step_fn(dir_path, batch_ops):
            new_dirs = []
            try:
                items = os.listdir(dir_path)
            except Exception:
                self._add_runtime_warning(f"无法访问目录：{dir_path}")
                return ("skipped", new_dirs)

            for item in items:
                full_path = os.path.join(dir_path, item)
                is_dir = os.path.isdir(full_path)
                current_path = full_path

                if find_text in item:
                    new_name = item.replace(find_text, replace_text)
                    if new_name != item:
                        matched_count["value"] += 1
                        new_path, warning_message = self._prepare_rename_target(dir_path, item, new_name, "项目")
                        if warning_message:
                            skipped_count["value"] += 1
                            self._add_runtime_warning(warning_message)
                        elif new_path is None:
                            skipped_count["value"] += 1
                        else:
                            try:
                                os.rename(full_path, new_path)
                                batch_ops.append({
                                    'type': 'rename',
                                    'details': {
                                        'old_path': full_path,
                                        'new_path': new_path,
                                        'operation_type': 'rename_all_items'
                                    }
                                })
                                renamed_count["value"] += 1
                                current_path = new_path
                            except Exception:
                                failed_count["value"] += 1
                                self._add_runtime_warning(f"重命名失败：{item}")
                                continue

                if is_dir:
                    new_dirs.append(current_path)

                self._process_pending_ui()

            return ("processed", new_dirs)

        def on_done(summary):
            if summary.get("batch_operations"):
                desc = "全部目录字符替换" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {
                    'description': desc,
                    'operation_type': 'rename_all_items',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()

            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：成功 {renamed_count['value']} 项，跳过 {skipped_count['value']} 项，失败 {failed_count['value']} 项。")

            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("项目", matched_count["value"], renamed_count["value"], skipped_count["value"], failed_count["value"]),
                refresh_mode="directories_only",
                warning_title="重命名提示",
                update_history=renamed_count["value"] > 0
            )

        self.task_runner.start([root_path], step_fn, on_done)

    def multi_rename_current_level(self):
        """批量重命名本级目录"""
        if not self._require_source_directory():
            return

        # 20260906 201500 新增：非勾选模式下提示用户先切换到支持勾选的操作模式
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            self._warn_check_mode_required()
            return

        directories = self._get_source_items("directory")
        if directories is None:
            return

        if not directories:
            self._show_info_message("提示", "当前目录下没有子目录!")
            return

        # 20260906 184352 改为先勾选再重命名：仅处理用户在目录树中勾选的本级目录（支持“全选/取消全选”），未勾选时给出操作步骤提示
        source_path = self.source_dir.get()
        source_norm = os.path.normcase(os.path.abspath(source_path))
        selected_dirs = []
        for path in self.get_checked_directory_paths():  # 20260906 184352 仅取源目录下一级（本级）的已勾选目录
            parent_norm = os.path.normcase(os.path.abspath(os.path.dirname(path)))
            if parent_norm == source_norm:
                selected_dirs.append(os.path.basename(path))
        if not selected_dirs:
            self._show_warning_message(
                "提示",
                "请先选择需要重命名的目录！\n\n"
                "操作步骤：\n"
                "1. 在“复制选定目录和文件”或“复制选定层级目录”模式下浏览源目录；\n"
                "2. 勾选源目录下一级需要重命名的目录（可逐个勾选，或点击下方的“全选”按钮）；\n"
                "3. 再点击“多维重命名本级目录名”按钮。"
            )
            return

        self.create_multi_rename_dialog(selected_dirs)

    def multi_rename_current_files(self):
        """批量重命名本级文件"""
        if not self._require_source_directory():
            return

        # 20260906 201500 新增：非勾选模式下提示用户先切换到支持勾选的操作模式
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            self._warn_check_mode_required()
            return

        files = self._get_source_items("file")
        if files is None:
            return

        if not files:
            self._show_info_message("提示", "当前目录下没有文件!")
            return

        # 20260906 184352 改为先勾选再重命名：仅处理用户在目录树中勾选的本级文件（支持“全选/取消全选”），未勾选时给出操作步骤提示
        source_path = self.source_dir.get()
        source_norm = os.path.normcase(os.path.abspath(source_path))
        selected_files = []
        for path in self.get_checked_file_paths():  # 20260906 184352 仅取源目录下一级（本级）的已勾选文件
            parent_norm = os.path.normcase(os.path.abspath(os.path.dirname(path)))
            if parent_norm == source_norm:
                selected_files.append(os.path.basename(path))
        if not selected_files:
            self._show_warning_message(
                "提示",
                "请先选择需要重命名的文件！\n\n"
                "操作步骤：\n"
                "1. 在“复制选定目录和文件”模式下浏览源目录；\n"
                "2. 勾选源目录下一级需要重命名的文件（可逐个勾选，或点击下方的“全选”按钮）；\n"
                "3. 再点击“多维重命名本级文件名”按钮。"
            )
            return

        self.create_multi_rename_file_dialog(selected_files)
        
    def create_multi_rename_dialog(self, directories):
        """创建批量重命名对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("多维重命名本级目录名")
        dialog.geometry("1000x300")
        dialog.resizable(False, False)

        # 设置对话框位置：水平居中，距上边50像素
        dialog.update_idletasks()  # 确保窗口尺寸已计算
        screen_width = dialog.winfo_screenwidth()
        dialog_width = 1000
        x = (screen_width - dialog_width) // 2
        y = 50
        dialog.geometry(f"{dialog_width}x300+{x}+{y}")

        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 20260906 184352 新增：显示已勾选的目录数量，明确本次重命名的作用范围
        ttk.Label(main_frame, text=f"已选择 {len(directories)} 个目录", foreground="#3366CC", font=("TkDefaultFont", 10, "bold")).pack(anchor='e', pady=(0, 5))

        # 项目名称标签
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(title_frame, text="项目名称:", font=("TkDefaultFont", 10, "bold")).grid(row=0, column=0, sticky='w')
        ttk.Label(title_frame, text="前序号").grid(row=0, column=1, padx=(10,30))
        ttk.Label(title_frame, text="前连接符").grid(row=0, column=2, padx=30)
        ttk.Label(title_frame, text="前缀字符").grid(row=0, column=3, padx=30)
        ttk.Label(title_frame, text="原名称").grid(row=0, column=4, padx=(50,30))
        ttk.Label(title_frame, text="后缀字符").grid(row=0, column=6, padx=40)
        ttk.Label(title_frame, text="后连接符").grid(row=0, column=7, padx=20)
        ttk.Label(title_frame, text="后序号").grid(row=0, column=8, padx=20)
        
        # 项目选择复选框
        select_frame = ttk.Frame(main_frame)
        select_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(select_frame, text="项目选择:", font=("TkDefaultFont", 10, "bold")).grid(row=0, column=0, sticky='w')
        
        controls = {
            "prefix_num_var": tk.BooleanVar(),
            "prefix_conn_var": tk.BooleanVar(),
            "prefix_text_var": tk.BooleanVar(),
            "suffix_text_var": tk.BooleanVar(),
            "suffix_conn_var": tk.BooleanVar(),
            "suffix_num_var": tk.BooleanVar(),
        }
        
        ttk.Checkbutton(select_frame, text="前序号", variable=controls["prefix_num_var"]).grid(row=0, column=1, padx=(10,20))
        ttk.Checkbutton(select_frame, text="前连接符", variable=controls["prefix_conn_var"]).grid(row=0, column=2, padx=(10,20))
        ttk.Checkbutton(select_frame, text="前缀字符", variable=controls["prefix_text_var"]).grid(row=0, column=3, padx=10)
        ttk.Checkbutton(select_frame, text="后缀字符", variable=controls["suffix_text_var"]).grid(row=0, column=6, padx=(180,20))
        ttk.Checkbutton(select_frame, text="后连接符", variable=controls["suffix_conn_var"]).grid(row=0, column=7, padx=20)
        ttk.Checkbutton(select_frame, text="后序号", variable=controls["suffix_num_var"]).grid(row=0, column=8, padx=10)
        
        # 目录新名输入框架
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(input_frame, text="目录新名:", font=("TkDefaultFont", 10, "bold")).grid(row=0, column=0, sticky='w')
        
        # 序号类型选项
        num_options = ["1", "A", "a", "Ⅰ"]
        controls["prefix_num_combo"] = ttk.Combobox(input_frame, values=num_options, width=5, state="readonly")
        controls["prefix_num_combo"].grid(row=0, column=1, padx=(10,20))
        controls["prefix_num_combo"].set("1")
        
        # 连接符选项
        conn_options = [".", "_", "-", "–", "—", "•"]
        controls["prefix_conn_combo"] = ttk.Combobox(input_frame, values=conn_options, width=5, state="readonly")
        controls["prefix_conn_combo"].grid(row=0, column=2, padx=10)
        controls["prefix_conn_combo"].set(".")
        
        # 前缀文本输入框
        controls["prefix_text_entry"] = ttk.Entry(input_frame, width=12)
        controls["prefix_text_entry"].grid(row=0, column=3, padx=(20,10))
        
        # 添加说明文字
        note_label = ttk.Label(input_frame, text="+ 目录原名称 +", font=("TkDefaultFont", 10))
        note_label.grid(row=0, column=4, padx=0)
        
        # 后缀文本输入框
        controls["suffix_text_entry"] = ttk.Entry(input_frame, width=12)
        controls["suffix_text_entry"].grid(row=0, column=6, padx=10)
        
        # 后连接符选项
        controls["suffix_conn_combo"] = ttk.Combobox(input_frame, values=conn_options, width=5, state="readonly")
        controls["suffix_conn_combo"].grid(row=0, column=7, padx=20)
        controls["suffix_conn_combo"].set(".")
        
        # 后序号类型选项
        controls["suffix_num_combo"] = ttk.Combobox(input_frame, values=num_options, width=5, state="readonly")
        controls["suffix_num_combo"].grid(row=0, column=8, padx=10)
        controls["suffix_num_combo"].set("1")
        
        # 按钮框架
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=20)
        
        def do_multi_rename():
            try:
                self.execute_multi_rename(directories, controls)
                dialog.destroy()
            except Exception as e:
                self._show_error_message("错误", f"重命名过程中出错: {str(e)}")
        
        ttk.Button(button_frame, text="确定", command=do_multi_rename).pack(side=tk.LEFT, padx=10)
        ttk.Button(button_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=10)
        
    def execute_multi_rename(self, directories, controls=None):
        """执行批量重命名"""
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，请稍后再试。")
            return  # 20260402 094929 避免并发任务导致状态混乱

        source_path = self.source_dir.get()
        self._reset_runtime_warnings()
        controls = self._resolve_multi_rename_controls(controls, "directory")
        plans = []
        for i, dir_name in enumerate(directories):
            new_name = self._build_multi_rename_name(dir_name, i + 1, controls)
            if new_name != dir_name:
                plans.append((dir_name, new_name))

        if not plans:
            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("目录", 0, 0, 0, 0),
                refresh_mode="directories_only",
                warning_title="重命名提示",
                update_history=False
            )
            return

        self.task_runner = self.TaskRunner(self, description="多维重命名本级目录名", total=len(plans))  # 20260402 094929 分片执行

        def step_fn(item, batch_ops):
            dir_name, new_name = item
            old_path = os.path.join(source_path, dir_name)
            new_path, warning_message = self._prepare_rename_target(source_path, dir_name, new_name, "目录")
            if warning_message:
                self._add_runtime_warning(warning_message)
                return "skipped"
            if new_path is None:
                return "skipped"
            try:
                os.rename(old_path, new_path)
                batch_ops.append({
                    'type': 'rename',
                    'details': {
                        'old_path': old_path,
                        'new_path': new_path,
                        'operation_type': 'multi_rename_directory'
                    }
                })
            except Exception:
                self._add_runtime_warning(f"重命名失败：{dir_name}")
                raise

        def on_done(summary):
            matched_count = len(plans)
            renamed_count = summary.get("completed", 0)
            skipped_count = summary.get("skipped", 0)
            failed_count = summary.get("failed", 0)

            if summary.get("batch_operations"):
                desc = "多维重命名本级目录名" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {
                    'description': desc,
                    'operation_type': 'multi_rename_directory',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()

            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {renamed_count} 项，跳过 {skipped_count} 项，失败 {failed_count} 项。")

            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("目录", matched_count, renamed_count, skipped_count, failed_count),
                refresh_mode="directories_only",
                warning_title="重命名提示",
                update_history=renamed_count > 0
            )

        self.task_runner.start(plans, step_fn, on_done)
        
    def create_multi_rename_file_dialog(self, files):
        """创建批量文件重命名对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("多维重命名本级文件名")
        dialog.geometry("1000x300")
        dialog.resizable(False, False)

        # 设置对话框位置：水平居中，距上边50像素
        dialog.update_idletasks()  # 确保窗口尺寸已计算
        screen_width = dialog.winfo_screenwidth()
        dialog_width = 1000
        x = (screen_width - dialog_width) // 2
        y = 50
        dialog.geometry(f"{dialog_width}x300+{x}+{y}")

        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 20260906 184352 新增：显示已勾选的文件数量，明确本次重命名的作用范围
        ttk.Label(main_frame, text=f"已选择 {len(files)} 个文件", foreground="#3366CC", font=("TkDefaultFont", 10, "bold")).pack(anchor='e', pady=(0, 5))

        # 项目名称标签
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(title_frame, text="项目名称:", font=("TkDefaultFont", 10, "bold")).grid(row=0, column=0, sticky='w')
        ttk.Label(title_frame, text="前序号").grid(row=0, column=1, padx=(10,30))
        ttk.Label(title_frame, text="前连接符").grid(row=0, column=2, padx=30)
        ttk.Label(title_frame, text="前缀字符").grid(row=0, column=3, padx=30)
        ttk.Label(title_frame, text="原名称").grid(row=0, column=4, padx=(50,30))
        ttk.Label(title_frame, text="后缀字符").grid(row=0, column=6, padx=40)
        ttk.Label(title_frame, text="后连接符").grid(row=0, column=7, padx=20)
        ttk.Label(title_frame, text="后序号").grid(row=0, column=8, padx=20)
        
        # 项目选择复选框
        select_frame = ttk.Frame(main_frame)
        select_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(select_frame, text="项目选择:", font=("TkDefaultFont", 10, "bold")).grid(row=0, column=0, sticky='w')
        
        controls = {
            "prefix_num_var": tk.BooleanVar(),
            "prefix_conn_var": tk.BooleanVar(),
            "prefix_text_var": tk.BooleanVar(),
            "suffix_text_var": tk.BooleanVar(),
            "suffix_conn_var": tk.BooleanVar(),
            "suffix_num_var": tk.BooleanVar(),
        }
        
        ttk.Checkbutton(select_frame, text="前序号", variable=controls["prefix_num_var"]).grid(row=0, column=1, padx=(10,20))
        ttk.Checkbutton(select_frame, text="前连接符", variable=controls["prefix_conn_var"]).grid(row=0, column=2, padx=(10,20))
        ttk.Checkbutton(select_frame, text="前缀字符", variable=controls["prefix_text_var"]).grid(row=0, column=3, padx=10)
        ttk.Checkbutton(select_frame, text="后缀字符", variable=controls["suffix_text_var"]).grid(row=0, column=6, padx=(180,20))
        ttk.Checkbutton(select_frame, text="后连接符", variable=controls["suffix_conn_var"]).grid(row=0, column=7, padx=20)
        ttk.Checkbutton(select_frame, text="后序号", variable=controls["suffix_num_var"]).grid(row=0, column=8, padx=10)
        
        # 文件新名输入框架
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(input_frame, text="文件新名:", font=("TkDefaultFont", 10, "bold")).grid(row=0, column=0, sticky='w')
        
        # 序号类型选项
        num_options = ["1", "A", "a", "Ⅰ"]
        controls["prefix_num_combo"] = ttk.Combobox(input_frame, values=num_options, width=5, state="readonly")
        controls["prefix_num_combo"].grid(row=0, column=1, padx=(10,20))
        controls["prefix_num_combo"].set("1")
        
        # 连接符选项
        conn_options = [".", "_", "-", "–", "—", "•"]
        controls["prefix_conn_combo"] = ttk.Combobox(input_frame, values=conn_options, width=5, state="readonly")
        controls["prefix_conn_combo"].grid(row=0, column=2, padx=10)
        controls["prefix_conn_combo"].set(".")
        
        # 前缀文本输入框
        controls["prefix_text_entry"] = ttk.Entry(input_frame, width=12)
        controls["prefix_text_entry"].grid(row=0, column=3, padx=(20,10))
        
        # 添加说明文字
        note_label = ttk.Label(input_frame, text="+ 文件原名称 +", font=("TkDefaultFont", 10))
        note_label.grid(row=0, column=4, padx=0)
        
        # 后缀文本输入框
        controls["suffix_text_entry"] = ttk.Entry(input_frame, width=12)
        controls["suffix_text_entry"].grid(row=0, column=6, padx=10)
        
        # 后连接符选项
        controls["suffix_conn_combo"] = ttk.Combobox(input_frame, values=conn_options, width=5, state="readonly")
        controls["suffix_conn_combo"].grid(row=0, column=7, padx=20)
        controls["suffix_conn_combo"].set(".")
        
        # 后序号类型选项
        controls["suffix_num_combo"] = ttk.Combobox(input_frame, values=num_options, width=5, state="readonly")
        controls["suffix_num_combo"].grid(row=0, column=8, padx=10)
        controls["suffix_num_combo"].set("1")
        
        # 按钮框架
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=20)
        
        def do_multi_file_rename():
            try:
                self.execute_multi_file_rename(files, controls)
                dialog.destroy()
            except Exception as e:
                self._show_error_message("错误", f"重命名过程中出错: {str(e)}")
        
        ttk.Button(button_frame, text="确定", command=do_multi_file_rename).pack(side=tk.LEFT, padx=10)
        ttk.Button(button_frame, text="取消", command=dialog.destroy).pack(side=tk.LEFT, padx=10)
        
    def execute_multi_file_rename(self, files, controls=None):
        """执行批量文件重命名"""
        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，请稍后再试。")
            return  # 20260402 094929 避免并发任务导致状态混乱

        source_path = self.source_dir.get()
        self._reset_runtime_warnings()
        controls = self._resolve_multi_rename_controls(controls, "file")
        plans = []
        for i, file_name in enumerate(files):
            name_part, ext_part = os.path.splitext(file_name)
            new_name = self._build_multi_rename_name(name_part, i + 1, controls, ext_part)
            if new_name != file_name:
                plans.append((file_name, new_name))

        if not plans:
            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("文件", 0, 0, 0, 0),
                refresh_mode="files_only",
                warning_title="重命名提示",
                update_history=False
            )
            return

        self.task_runner = self.TaskRunner(self, description="多维重命名本级文件名", total=len(plans))  # 20260402 094929 分片执行

        def step_fn(item, batch_ops):
            file_name, new_name = item
            old_path = os.path.join(source_path, file_name)
            new_path, warning_message = self._prepare_rename_target(source_path, file_name, new_name, "文件")
            if warning_message:
                self._add_runtime_warning(warning_message)
                return "skipped"
            if new_path is None:
                return "skipped"
            try:
                os.rename(old_path, new_path)
                batch_ops.append({
                    'type': 'rename',
                    'details': {
                        'old_path': old_path,
                        'new_path': new_path,
                        'operation_type': 'multi_rename_file'
                    }
                })
            except Exception:
                self._add_runtime_warning(f"重命名失败：{file_name}")
                raise

        def on_done(summary):
            matched_count = len(plans)
            renamed_count = summary.get("completed", 0)
            skipped_count = summary.get("skipped", 0)
            failed_count = summary.get("failed", 0)

            if summary.get("batch_operations"):
                desc = "多维重命名本级文件名" + ("（已取消）" if summary.get("cancelled") else "")
                self.operation_history.add_operation('batch', {
                    'description': desc,
                    'operation_type': 'multi_rename_file',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()

            if summary.get("cancelled"):
                self._show_info_message("提示", f"已取消：完成 {renamed_count} 项，跳过 {skipped_count} 项，失败 {failed_count} 项。")

            self._finalize_tree_change(
                "完成",
                self._build_rename_summary_message("文件", matched_count, renamed_count, skipped_count, failed_count),
                refresh_mode="files_only",
                warning_title="重命名提示",
                update_history=renamed_count > 0
            )

        self.task_runner.start(plans, step_fn, on_done)
        
    def generate_sequence_number(self, index, num_type):
        """生成序列号"""
        if num_type == "1":
            return str(index)
        elif num_type == "A":
            # 大写字母序列 A, B, C, ..., Z, AA, AB, ...
            result = ""
            while index > 0:
                index -= 1
                result = chr(65 + index % 26) + result
                index //= 26
            return result
        elif num_type == "a":
            # 小写字母序列 a, b, c, ..., z, aa, ab, ...
            result = ""
            while index > 0:
                index -= 1
                result = chr(97 + index % 26) + result
                index //= 26
            return result
        elif num_type == "Ⅰ":
            # 罗马数字
            roman_numerals = [
                (1000, 'Ⅿ'), (900, 'ⅭⅯ'), (500, 'Ⅾ'), (400, 'ⅭⅮ'),
                (100, 'Ⅽ'), (90, 'ⅩⅭ'), (50, 'Ⅼ'), (40, 'ⅩⅬ'),
                (10, 'Ⅹ'), (9, 'Ⅸ'), (5, 'Ⅴ'), (4, 'Ⅳ'), (1, 'Ⅰ')
            ]
            result = ""
            for value, numeral in roman_numerals:
                count = index // value
                result += numeral * count
                index -= value * count
            return result
        else:
            return str(index)
    
    def get_size_in_kb(self, path):
        """获取文件或目录的大小，单位为KB"""
        try:
            if os.path.isfile(path):
                # 文件大小
                size_bytes = os.path.getsize(path)
                return self._format_size_in_kb(size_bytes)
            elif os.path.isdir(path):
                total_size = self.directory_size_cache.get(path)
                if total_size is None:
                    return "N/A"
                return self._format_size_in_kb(total_size)
            else:
                return "N/A"
        except (OSError, IOError):
            return "N/A"
    
    def advanced_rename_directories(self):
        """全层级目录名称修改 - 基于已勾选目录，支持精确匹配、通配符和正则表达式 20260906 184352 改造"""
        if not self._require_source_directory():
            return

        # 20260906 201500 新增：非勾选模式下提示用户先切换到支持勾选的操作模式
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            self._warn_check_mode_required()
            return

        # 20260906 184352 改为先勾选再修改：不再调用 refresh_tree()（避免清空用户已勾选的目录）
        selected_dirs = self.get_checked_directory_paths()
        if not selected_dirs:
            self._show_warning_message(
                "提示",
                "请先选择需要修改名称的目录！\n\n"
                "操作步骤：\n"
                "1. 在“复制选定目录和文件”或“复制选定层级目录”模式下浏览源目录；\n"
                "2. 勾选需要修改名称的目录（可展开目录逐层勾选，可逐个勾选或点击下方的“全选”按钮，也可用“取消全选”重新选择）；\n"
                "3. 再点击“全层级目录名称修改”按钮。"
            )
            return

        # 创建目录名称修改对话框（仅处理已勾选的目录）
        self.create_file_rename_dialog(selected_dirs, "directory")
    
    def advanced_rename_files(self):
        """各层级文件名称修改 - 基于已勾选文件（各层级），支持精确匹配、通配符和正则表达式 20260906 184352 更名"""
        if not self._require_source_directory():
            return

        # 20260906 201500 新增：非勾选模式下提示用户先切换到支持勾选的操作模式
        if self.copy_mode.get() not in ["custom", "selected_levels"]:
            self._warn_check_mode_required()
            return

        # 20260828 104445 不再刷新目录树，避免清空用户已勾选的文件
        selected_files = self.get_checked_file_paths()
        if not selected_files:
            self._show_warning_message(
                "提示",
                "请先选择需要修改文件名的文件！\n\n"
                "操作步骤：\n"
                "1. 在“复制选定目录和文件”模式下浏览源目录；\n"
                "2. 逐个勾选需要修改文件名的文件，或点击下方的“全选”按钮；\n"
                "3. 再点击“各层级文件名称修改”按钮。"
            )
            return

        # 创建文件名修改对话框（仅处理已勾选的文件）
        self.create_file_rename_dialog(selected_files, "file")
    
    def get_checked_file_paths(self):
        """获取当前已勾选的文件完整路径列表（仅文件，不含目录）"""
        files = []
        for item_id in list(self.checked_items):
            path = self.tree_item_paths.get(item_id)
            if path and os.path.isfile(path):
                files.append(path)
        # 20260828 104445 去重并按路径排序，保证处理顺序稳定
        return sorted(set(files), key=lambda p: p.lower())

    def get_checked_directory_paths(self):
        """获取当前已勾选的目录完整路径列表（仅目录，不含文件）20260906 184352 新增"""
        dirs = []
        for item_id in list(self.checked_items):
            path = self.tree_item_paths.get(item_id)
            if path and os.path.isdir(path):
                dirs.append(path)
        # 20260906 184352 去重并按路径排序，保证处理顺序稳定
        return sorted(set(dirs), key=lambda p: p.lower())
    
    def _make_matcher(self, find_pattern, replace_text, mode, case_sensitive, glob_replace=False):
        """根据匹配模式创建匹配函数，返回 match_item(name)->(matched, new_name)
        glob_replace=True（名称修改对话框使用，20260906 184352 起目录与文件对话框均使用）时，通配符替换为直观语义：
        替换为中的 * 和 ? 表示原名称中对应位置匹配到的内容，例如 *.txt 替换为 *.md 可将 a.txt 改为 a.md"""
        import re
        import fnmatch
        if mode == "wildcard" and glob_replace:
            # 20260828 104445 通配符直观替换：将查找通配符转成捕获组，替换为中的 * / ? 按位置取回捕获内容
            flags = (0 if case_sensitive else re.IGNORECASE) | re.DOTALL
            regex_parts = []
            for ch in find_pattern:
                if ch == '*':
                    regex_parts.append('(.*)')
                elif ch == '?':
                    regex_parts.append('(.)')
                else:
                    regex_parts.append(re.escape(ch))
            wildcard_regex = re.compile(''.join(regex_parts) + r'\Z', flags)

            def match_item(name):
                m = wildcard_regex.fullmatch(name)
                if m is None:
                    return False, name
                groups = m.groups()
                parts = []
                gi = 0
                for ch in replace_text:
                    if ch in ('*', '?'):
                        if gi < len(groups):
                            parts.append(groups[gi])
                        gi += 1
                    else:
                        parts.append(ch)
                return True, ''.join(parts)
        elif mode == "exact":
            if case_sensitive:
                def match_item(name):
                    matched = name == find_pattern
                    return matched, replace_text if matched else name
            else:
                pattern_lower = find_pattern.lower()

                def match_item(name):
                    matched = name.lower() == pattern_lower
                    return matched, replace_text if matched else name
        elif mode == "wildcard":
            flags = 0 if case_sensitive else re.IGNORECASE
            wildcard_regex = re.compile(fnmatch.translate(find_pattern), flags)

            def match_item(name):
                matched = wildcard_regex.fullmatch(name) is not None
                return matched, wildcard_regex.sub(replace_text, name, count=1) if matched else name
        elif mode == "regex":
            flags = 0 if case_sensitive else re.IGNORECASE
            try:
                regex_pattern = re.compile(find_pattern, flags)
            except re.error as e:
                raise Exception(f"正则表达式错误: {str(e)}")

            def match_item(name):
                matched = regex_pattern.search(name) is not None
                return matched, regex_pattern.sub(replace_text, name) if matched else name
        else:
            def match_item(name):
                return False, name
        return match_item

    def create_file_rename_dialog(self, selected_paths, target_type="file"):
        """创建名称修改对话框（仅针对已勾选的文件/目录，提供实时匹配数量提示）
        20260906 184352 扩展支持目录类型：target_type="file" 为“各层级文件名称修改”，"directory" 为“全层级目录名称修改”"""
        type_name = "文件" if target_type == "file" else "目录"
        title = "各层级文件名称修改" if target_type == "file" else "全层级目录名称修改"
        dialog = tk.Toplevel(self.root)
        dialog.title(title)

        # 设置对话框的最小尺寸，确保按钮始终可见
        dialog.minsize(600, 460)
        dialog.resizable(True, True)
        dialog.transient(self.root)
        dialog.grab_set()

        dialog.grid_rowconfigure(0, weight=1)
        dialog.grid_columnconfigure(0, weight=1)

        main_frame = ttk.Frame(dialog, padding="15")
        main_frame.grid(row=0, column=0, sticky='nsew')
        main_frame.grid_columnconfigure(0, weight=1)

        # 标题与已选数量（20260906 184352 支持文件/目录两种类型）
        title_frame = ttk.Frame(main_frame)
        title_frame.grid(row=0, column=0, sticky='ew', pady=(0, 10))
        ttk.Label(title_frame, text=title, font=("TkDefaultFont", 12, "bold")).pack(side=tk.LEFT)
        ttk.Label(title_frame, text=f"已选择 {len(selected_paths)} 个{type_name}", foreground="#3366CC").pack(side=tk.RIGHT)

        # 匹配模式
        match_mode = tk.StringVar(value="exact")
        mode_frame = ttk.LabelFrame(main_frame, text="匹配模式", padding="10")
        mode_frame.grid(row=1, column=0, sticky='ew', pady=(0, 10))

        ttk.Radiobutton(mode_frame, text="精确匹配", value="exact", variable=match_mode).grid(row=0, column=0, sticky='w', padx=(0, 20))
        ttk.Radiobutton(mode_frame, text="通配符匹配", value="wildcard", variable=match_mode).grid(row=0, column=1, sticky='w', padx=(0, 20))
        ttk.Radiobutton(mode_frame, text="正则表达式", value="regex", variable=match_mode).grid(row=0, column=2, sticky='w')

        mode_desc_var = tk.StringVar()
        ttk.Label(mode_frame, textvariable=mode_desc_var, font=("TkDefaultFont", 9), foreground="#666666", wraplength=560).grid(row=1, column=0, columnspan=3, sticky='w', pady=(6, 0))

        # 20260906 184352 说明文字按文件/目录类型分别给出示例
        if target_type == "file":
            mode_descriptions = {
                "exact": "说明：仅当文件名（含扩展名）与“查找内容”完全一致时才进行替换。",
                "wildcard": "说明：按整个文件名匹配，支持 *（任意多个字符）和 ?（单个字符）。替换为中的 * 和 ? 表示原文件名中对应位置的内容，例如 *.txt 替换为 *.md，可将 a.txt 改为 a.md。",
                "regex": r"说明：按正则表达式匹配，支持字符类、分组等高级语法，例如 ^IMG_\d+ 可匹配 IMG_1、IMG_2 等。",
            }
        else:
            mode_descriptions = {
                "exact": "说明：仅当目录名称与“查找内容”完全一致时才进行替换。",
                "wildcard": "说明：按整个目录名称匹配，支持 *（任意多个字符）和 ?（单个字符）。替换为中的 * 和 ? 表示原目录名中对应位置的内容，例如 项目* 替换为 归档_*，可将 项目2024 改为 归档_2024。",
                "regex": r"说明：按正则表达式匹配，支持字符类、分组等高级语法，例如 ^第\d+章 可匹配 第1章、第2章 等。",
            }

        # 查找和替换输入
        input_frame = ttk.LabelFrame(main_frame, text="查找和替换", padding="10")
        input_frame.grid(row=2, column=0, sticky='ew', pady=(0, 10))

        find_var = tk.StringVar()
        replace_var = tk.StringVar()

        ttk.Label(input_frame, text="查找内容:").grid(row=0, column=0, sticky='w', pady=(0, 5))
        find_entry = ttk.Entry(input_frame, textvariable=find_var, width=50)
        find_entry.grid(row=1, column=0, sticky='ew', pady=(0, 5))

        ttk.Label(input_frame, text="替换为:").grid(row=2, column=0, sticky='w', pady=(0, 5))
        replace_entry = ttk.Entry(input_frame, textvariable=replace_var, width=50)
        replace_entry.grid(row=3, column=0, sticky='ew')

        match_count_var = tk.StringVar(value="请输入查找内容后查看匹配数量")
        ttk.Label(input_frame, textvariable=match_count_var, font=("TkDefaultFont", 9, "bold"), foreground="#CC6600").grid(row=4, column=0, sticky='w', pady=(8, 0))

        input_frame.grid_columnconfigure(0, weight=1)

        # 选项
        options_frame = ttk.LabelFrame(main_frame, text="选项", padding="10")
        options_frame.grid(row=3, column=0, sticky='ew', pady=(0, 5))

        case_sensitive = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="区分大小写", variable=case_sensitive).grid(row=0, column=0, sticky='w', pady=(0, 5))

        preview_mode = tk.BooleanVar(value=True)
        ttk.Checkbutton(options_frame, text="预览模式（推荐：先预览将要修改的项目，确认后再执行）", variable=preview_mode).grid(row=1, column=0, sticky='w')

        # 按钮框架 - 固定在底部
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, sticky='ew', pady=(12, 0))

        def do_file_rename():
            find_pattern = find_var.get()
            replace_text = replace_var.get()
            mode = match_mode.get()
            case_sens = case_sensitive.get()
            preview = preview_mode.get()

            if not find_pattern:
                self._show_warning_message("警告", "请输入查找内容！")
                return

            try:
                # 20260906 184352 目录与文件统一走已选路径模式，通配符使用直观替换语义
                # 20260906 205815 传入 on_preview_confirm 回调，供预览对话框"同意预览结果执行修改"按钮直接触发执行
                confirm_cb = None
                if preview:
                    def confirm_cb():
                        preview_mode.set(False)
                        do_file_rename()
                self.execute_advanced_rename(find_pattern, replace_text, mode, case_sens, preview, target_type, selected_paths, True, on_preview_confirm=confirm_cb)
                if not preview:
                    dialog.destroy()
                    # 20260828 104445 执行完成后由任务结束回调按当前显示筛选刷新目录树以显示新文件名
            except Exception as e:
                self._show_error_message("错误", f"操作过程中出错:\n{str(e)}")

        # 实时匹配数量统计（防抖）
        self._file_rename_after_id = None

        def count_matches():
            find_text = find_var.get()
            if not find_text:
                match_count_var.set("请输入查找内容后查看匹配数量")
                return
            try:
                matcher = self._make_matcher(find_text, replace_var.get(), match_mode.get(), case_sensitive.get(), True)
            except Exception as e:
                match_count_var.set(f"表达式错误：{e}")
                return
            # 20260906 184352 按目标类型（文件/目录）统计匹配数量
            check_fn = os.path.isfile if target_type == "file" else os.path.isdir
            count = 0
            for path in selected_paths:
                if check_fn(path) and matcher(os.path.basename(path))[0]:
                    count += 1
            match_count_var.set(f"按当前条件匹配到 {count} 个{type_name}")

        def schedule_count_update(*args):
            if self._file_rename_after_id is not None:
                try:
                    self.root.after_cancel(self._file_rename_after_id)
                except Exception:
                    pass
            self._file_rename_after_id = self.root.after(150, count_matches)

        def update_mode_description(*args):
            mode_desc_var.set(mode_descriptions.get(match_mode.get(), ""))

        find_var.trace_add("write", schedule_count_update)
        replace_var.trace_add("write", schedule_count_update)
        case_sensitive.trace_add("write", schedule_count_update)
        match_mode.trace_add("write", lambda *a: (update_mode_description(), schedule_count_update()))
        update_mode_description()
        count_matches()

        # 创建按钮并居中显示
        button_container = ttk.Frame(button_frame)
        button_container.grid(row=0, column=0)
        button_frame.grid_columnconfigure(0, weight=1)  # 让按钮容器居中

        ttk.Button(button_container, text="执行修改", command=do_file_rename).grid(row=0, column=0, padx=(0, 10))
        ttk.Button(button_container, text="取消", command=dialog.destroy).grid(row=0, column=1)

        # 动态计算并设置对话框尺寸
        dialog.update_idletasks()
        req_width = main_frame.winfo_reqwidth() + 30
        req_height = main_frame.winfo_reqheight() + 50
        dialog_width = max(600, req_width)
        dialog_height = max(460, req_height)
        screen_width = dialog.winfo_screenwidth()
        screen_height = dialog.winfo_screenheight()
        max_height = int(screen_height * 0.8)
        dialog_height = min(dialog_height, max_height)
        x = (screen_width - dialog_width) // 2
        y = 50
        if y + dialog_height > screen_height - 50:
            y = max(50, screen_height - dialog_height - 50)
        dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

        find_entry.focus_set()

    def _export_rename_records(self, batch_operations, type_name):  # 20260906 205815 新增：导出重命名记录到 CSV/TXT
        """将本次重命名成功记录导出为 CSV 或 TXT 文件"""
        default_name = f"重命名记录_{self.get_current_time().replace(':', '').replace(' ', '_')}"
        save_path = filedialog.asksaveasfilename(
            title="导出重命名记录",
            defaultextension=".csv",
            initialfile=default_name,
            filetypes=[("CSV 文件", "*.csv"), ("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if not save_path:
            return

        try:
            is_csv = save_path.lower().endswith(".csv")
            sep = "," if is_csv else "\t"
            lines = [f"序号{sep}类型{sep}原路径{sep}新路径"]
            for i, op in enumerate(batch_operations, 1):
                details = op.get("details", {})
                old_path = details.get("old_path", "")
                new_path = details.get("new_path", "")
                op_type = details.get("operation_type", "")
                item_type = type_name if op_type else "项目"
                if is_csv:
                    # CSV 中路径含逗号时需用双引号包裹
                    old_path = f'"{old_path}"'
                    new_path = f'"{new_path}"'
                lines.append(f"{i}{sep}{item_type}{sep}{old_path}{sep}{new_path}")

            encoding = "utf-8-sig" if is_csv else "utf-8"  # CSV 用带 BOM 编码以便 Excel 正确识别中文
            with open(save_path, "w", encoding=encoding) as f:
                f.write("\n".join(lines))
            self._show_info_message("完成", f"已导出 {len(batch_operations)} 条重命名记录到:\n{save_path}")
        except Exception as e:
            self._show_error_message("错误", f"导出重命名记录失败:\n{str(e)}")

    def execute_advanced_rename(self, find_pattern, replace_text, mode, case_sensitive, preview_mode, target_type, selected_paths=None, glob_replace=False, on_preview_confirm=None):
        """执行高级重命名操作"""
        # 20260828 104445 新增 selected_paths 参数：名称修改时仅处理已勾选的文件/目录；为 None 时保持原有递归遍历全部项目的行为（兼容原测试用例）
        # 20260828 104445 新增 glob_replace 参数：名称修改对话框的通配符模式使用直观替换语义
        # 20260906 184352 selected_paths 同时支持文件与目录两种类型（各层级文件名称修改 / 全层级目录名称修改）
        # 20260906 205815 新增 on_preview_confirm 参数：预览模式下传入确认回调，供预览对话框"同意预览结果执行修改"按钮直接触发执行

        renamed_items = []
        preview_items = []
        self._reset_runtime_warnings()
        preview_virtual_names = {}
        batch_operations = []  # 20260402 085047 将单次高级重命名记录为批次操作，便于一次撤销/重做

        # 20260828 104445 匹配函数抽取为公共方法 _make_matcher，供对话框实时统计与执行共用
        match_item = self._make_matcher(find_pattern, replace_text, mode, case_sensitive, glob_replace)
        selected_files_mode = (target_type == "file" and selected_paths is not None)
        selected_dirs_mode = (target_type == "directory" and selected_paths is not None)  # 20260906 184352 新增已选目录模式
        selected_mode = selected_files_mode or selected_dirs_mode  # 20260906 184352 已选路径模式（文件或目录）
        def process_directory(path):
            # 20260828 111834 清理死代码：本函数仅在预览模式下调用，原非预览分支与冗余判断已移除
            try:
                items = sorted(os.listdir(path), key=lambda name: name.lower())
                preview_virtual_names[path] = set(items)
                for item in items:
                    full_path = os.path.join(path, item)
                    is_dir = os.path.isdir(full_path)

                    # 根据目标类型过滤
                    if target_type == "directory" and not is_dir:
                        continue
                    elif target_type == "file" and is_dir:
                        # 对于文件模式，仍需要递归处理目录
                        process_directory(full_path)
                        continue

                    match_found, new_name = match_item(item)

                    if match_found and new_name != item:
                        item_label = "目录" if is_dir else "文件"
                        try:
                            relative_path = os.path.relpath(full_path, self.source_dir.get())
                        except Exception:
                            relative_path = full_path
                        virtual_names = preview_virtual_names.get(path, set())
                        _, warning_message = self._prepare_rename_target(path, item, new_name, item_label)
                        if warning_message is None and new_name in virtual_names and new_name != item:
                            warning_message = f"{item_label} '{new_name}' 在本批次执行中会冲突，跳过重命名 '{item}'"
                        if warning_message is None:
                            if item in virtual_names:
                                virtual_names.remove(item)
                            virtual_names.add(new_name)
                            preview_virtual_names[path] = virtual_names
                        preview_items.append({
                            'path': full_path,
                            'relative_path': relative_path,
                            'old_name': item,
                            'new_name': new_name,
                            'type': item_label,
                            'status': warning_message or "可执行"
                        })

                    # 递归处理子目录
                    if is_dir:
                        process_directory(full_path)
                    self._process_pending_ui()

            except (PermissionError, OSError) as e:
                self._add_runtime_warning(f"无法访问目录：{path}")
        
        if preview_mode:
            if selected_mode:
                # 20260828 104445 已选文件预览：仅遍历勾选的文件，不递归目录
                # 20260906 184352 已选目录预览：仅遍历勾选的目录，不递归其子级（逻辑与已选文件一致）
                for full_path in selected_paths:
                    is_dir = os.path.isdir(full_path)
                    if target_type == "file" and not os.path.isfile(full_path):
                        continue
                    if target_type == "directory" and not is_dir:
                        continue
                    item = os.path.basename(full_path)
                    match_found, new_name = match_item(item)
                    if match_found and new_name != item:
                        item_label = "目录" if is_dir else "文件"
                        try:
                            relative_path = os.path.relpath(full_path, self.source_dir.get())
                        except Exception:
                            relative_path = full_path
                        parent_dir = os.path.dirname(full_path)
                        virtual_names = preview_virtual_names.get(parent_dir, set())
                        _, warning_message = self._prepare_rename_target(parent_dir, item, new_name, item_label)
                        if warning_message is None and new_name in virtual_names and new_name != item:
                            warning_message = f"{item_label} '{new_name}' 在本批次执行中会冲突，跳过重命名 '{item}'"
                        if warning_message is None:
                            if item in virtual_names:
                                virtual_names.remove(item)
                            virtual_names.add(new_name)
                            preview_virtual_names[parent_dir] = virtual_names
                        preview_items.append({
                            'path': full_path,
                            'relative_path': relative_path,
                            'old_name': item,
                            'new_name': new_name,
                            'type': item_label,
                            'status': warning_message or "可执行"
                        })
                    self._process_pending_ui()
            else:
                process_directory(self.source_dir.get())
            if preview_items:
                self.show_preview_dialog(preview_items, on_confirm=on_preview_confirm)
            else:
                self._show_info_message("预览结果", "没有找到匹配的项目。")
            return

        if getattr(self, "task_running", False):
            self._show_warning_message("提示", "当前有任务正在执行，请稍后再试。")
            return  # 20260402 094929 避免并发任务导致状态混乱

        counters = {"matched": 0, "renamed": 0, "skipped": 0, "failed": 0}  # 20260402 094929 执行模式计数器
        self.task_runner = self.TaskRunner(self, description="高级重命名")  # 20260402 094929 分片执行（总量不预估）

        def step_fn(dir_path, batch_ops):
            if selected_mode:
                # 20260828 104445 已选文件模式：dir_path 为单个文件路径，仅处理该文件，不递归
                # 20260906 184352 已选目录模式：dir_path 为单个目录路径，仅处理该目录，不递归（任务队列已按层级深度降序排列，先深层后浅层，避免先改父目录导致子目录路径失效）
                try:
                    full_path = dir_path
                    is_dir = os.path.isdir(full_path)
                    if target_type == "file" and not os.path.isfile(full_path):
                        return ("processed", [])
                    if target_type == "directory" and not is_dir:
                        return ("processed", [])
                    item = os.path.basename(full_path)
                    match_found, new_name = match_item(item)
                    if match_found and new_name != item:
                        counters["matched"] += 1
                        item_label = "目录" if is_dir else "文件"
                        new_path, warning_message = self._prepare_rename_target(os.path.dirname(full_path), item, new_name, item_label)
                        if warning_message:
                            counters["skipped"] += 1
                            self._add_runtime_warning(warning_message)
                        elif new_path is None:
                            counters["skipped"] += 1
                        else:
                            try:
                                os.rename(full_path, new_path)
                                batch_ops.append({
                                    'type': 'rename',
                                    'details': {
                                        'old_path': full_path,
                                        'new_path': new_path,
                                        'operation_type': f'advanced_rename_{target_type}'
                                    }
                                })
                                counters["renamed"] += 1
                            except Exception:
                                counters["failed"] += 1
                                self._add_runtime_warning(f"重命名失败：{item}")
                    self._process_pending_ui()
                except (PermissionError, OSError):
                    self._add_runtime_warning(f"无法访问：{dir_path}")
                return ("processed", [])
            new_dirs = []
            try:
                items = sorted(os.listdir(dir_path), key=lambda name: name.lower())
            except (PermissionError, OSError):
                self._add_runtime_warning(f"无法访问目录：{dir_path}")
                return ("skipped", new_dirs)

            for item in items:
                full_path = os.path.join(dir_path, item)
                is_dir = os.path.isdir(full_path)

                if target_type == "directory" and not is_dir:
                    continue

                if target_type == "file" and is_dir:
                    new_dirs.append(full_path)
                    continue

                match_found, new_name = match_item(item)
                if match_found and new_name != item:
                    counters["matched"] += 1
                    item_label = "目录" if is_dir else "文件"
                    new_path, warning_message = self._prepare_rename_target(dir_path, item, new_name, item_label)
                    if warning_message:
                        counters["skipped"] += 1
                        self._add_runtime_warning(warning_message)
                    elif new_path is None:
                        counters["skipped"] += 1
                    else:
                        try:
                            os.rename(full_path, new_path)
                            batch_ops.append({
                                'type': 'rename',
                                'details': {
                                    'old_path': full_path,
                                    'new_path': new_path,
                                    'operation_type': f'advanced_rename_{target_type}'
                                }
                            })
                            counters["renamed"] += 1
                            full_path = new_path
                        except Exception:
                            counters["failed"] += 1
                            self._add_runtime_warning(f"重命名失败：{item}")

                if is_dir:
                    new_dirs.append(full_path)

                self._process_pending_ui()

            return ("processed", new_dirs)

        def on_done(summary):
            if summary.get("batch_operations"):
                # 20260828 104445 已选文件模式的操作记录名称调整为“文件名称修改”
                # 20260906 184352 操作记录名称同步更新：已选目录→“全层级目录名称修改”，已选文件→“各层级文件名称修改”；selected_paths 为 None 的递归遍历旧路径保持原名称（兼容原测试用例）
                if target_type == "directory":
                    description = "全层级目录名称修改" if selected_dirs_mode else "全部目录名修改"
                elif selected_files_mode:
                    description = "各层级文件名称修改"
                else:
                    description = "全部文件名修改"
                if summary.get("cancelled"):
                    description += "（已取消）"
                self.operation_history.add_operation('batch', {
                    'description': description,
                    'operation_type': f'advanced_rename_{target_type}',
                    'operations': summary["batch_operations"]
                })
                self.update_history_buttons()

            type_name = "目录" if target_type == "directory" else "文件"
            message_prefix = "已取消！" if summary.get("cancelled") else "重命名完成！"
            self._finalize_tree_change(
                "完成",
                f"{message_prefix}匹配 {counters['matched']} 个{type_name}，成功 {counters['renamed']} 个，跳过 {counters['skipped']} 个，失败 {counters['failed']} 个。",  # 20260402 094929 执行结果汇总
                # 20260828 104445 已选文件模式按用户当前显示筛选刷新，不强制切换视图
                # 20260906 184352 已选目录模式同样按用户当前显示筛选刷新，不强制切换视图
                refresh_mode=(self.tree_display_mode.get() if selected_mode else ("directories_only" if target_type == "directory" else "files_only")),  # 20260402 094929 执行结束后刷新视图
                warning_title="重命名提示",
                update_history=counters["renamed"] > 0
            )
            # 20260906 205815 新增：重命名完成后询问是否导出重命名记录（CSV/TXT）
            if counters["renamed"] > 0 and summary.get("batch_operations"):
                if messagebox.askyesno("导出重命名记录", f"本次成功重命名 {counters['renamed']} 个{type_name}。\n是否导出重命名记录？\n（可导出为 CSV 或 TXT 格式）"):
                    self._export_rename_records(summary["batch_operations"], type_name)

        # 20260828 104445 已选文件模式以勾选文件列表作为任务队列；否则以源目录为起点递归遍历
        # 20260906 184352 已选目录模式以勾选目录列表作为任务队列，并按路径层级深度降序排列（先深层后浅层）：同时勾选父子目录时先改子目录再改父目录，避免父目录改名后子目录路径失效
        if selected_mode:
            initial_queue = list(selected_paths) if selected_files_mode else sorted(selected_paths, key=lambda p: p.count(os.sep), reverse=True)
        else:
            initial_queue = [self.source_dir.get()]
        self.task_runner.start(initial_queue, step_fn, on_done)  # 20260402 094929 启动高级重命名分片任务
    
    def show_preview_dialog(self, preview_items, on_confirm=None):
        """显示预览对话框

        Args:
            preview_items: 预览结果列表
            on_confirm: 可选回调，传入后在预览对话框底部显示"同意预览结果执行修改"按钮，
                        点击后先关闭预览对话框再调用此回调（用于跳过预览直接执行修改）  # 20260906 205815 新增
        """
        preview_dialog = tk.Toplevel(self.root)
        preview_dialog.title("预览重命名结果")
        preview_dialog.geometry("700x400")
        
        # 设置对话框位置
        preview_dialog.update_idletasks()
        screen_width = preview_dialog.winfo_screenwidth()
        dialog_width = 700
        x = (screen_width - dialog_width) // 2
        y = 80
        preview_dialog.geometry(f"{dialog_width}x400+{x}+{y}")
        
        preview_dialog.transient(self.root)
        preview_dialog.grab_set()
        
        main_frame = ttk.Frame(preview_dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        executable_count = sum(1 for item in preview_items if item.get('status') == "可执行")
        skipped_count = len(preview_items) - executable_count
        title_base_text = f"预览重命名结果 (共 {len(preview_items)} 个，可执行 {executable_count} 个，需跳过 {skipped_count} 个)"
        title_var = tk.StringVar()
        title_label = ttk.Label(main_frame, textvariable=title_var, 
                               font=("TkDefaultFont", 11, "bold"))
        title_label.pack(pady=(0, 10))

        show_only_skipped = tk.BooleanVar(value=self.preview_filter_mode == "skipped")
        show_only_executable = tk.BooleanVar(value=self.preview_filter_mode == "executable")
        filter_frame = ttk.Frame(main_frame)
        filter_frame.pack(fill=tk.X, pady=(0, 6))
        ttk.Checkbutton(
            filter_frame,
            text="仅看可执行项",
            variable=show_only_executable,
            command=lambda: on_filter_change("executable")
        ).pack(side=tk.LEFT)
        ttk.Checkbutton(
            filter_frame,
            text="仅看需跳过项",
            variable=show_only_skipped,
            command=lambda: on_filter_change("skipped")
        ).pack(side=tk.LEFT, padx=(8, 0))
        ttk.Button(filter_frame, text="重置筛选", command=lambda: reset_filters()).pack(side=tk.LEFT, padx=(12, 0))
        preview_summary_var = tk.StringVar()
        ttk.Label(filter_frame, textvariable=preview_summary_var, foreground="#666666").pack(side=tk.RIGHT)
        
        # 创建树形视图显示预览结果
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)
        
        preview_tree = ttk.Treeview(tree_frame, columns=("type", "old_name", "new_name", "status", "relative_path"), show="headings")
        preview_tree.heading("type", text="类型")
        preview_tree.heading("old_name", text="原名称")
        preview_tree.heading("new_name", text="新名称")
        preview_tree.heading("status", text="执行结果")
        preview_tree.heading("relative_path", text="路径")
        
        preview_tree.column("type", width=80)
        preview_tree.column("old_name", width=150)
        preview_tree.column("new_name", width=150)
        preview_tree.column("status", width=200)
        preview_tree.column("relative_path", width=320)
        
        # 添加滚动条
        v_scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=preview_tree.yview)
        preview_tree.configure(yscrollcommand=v_scrollbar.set)
        
        preview_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        preview_tree.tag_configure("skipped", foreground="#AA0000")
        
        # 填充数据
        sorted_preview_items = sorted(
            preview_items,
            key=lambda item: (
                0 if item.get('status', '可执行') == "可执行" else 1,
                str(item.get('relative_path', item.get('path', ''))).lower(),
                str(item.get('old_name', '')).lower()
            )
        )
        visible_items_state = {"items": []}
        button_refs = {"copy": None, "reason": None, "export": None}

        def on_filter_change(active_filter):
            if active_filter == "skipped" and show_only_skipped.get():
                show_only_executable.set(False)
                self.preview_filter_mode = "skipped"
            elif active_filter == "executable" and show_only_executable.get():
                show_only_skipped.set(False)
                self.preview_filter_mode = "executable"
            elif show_only_executable.get():
                self.preview_filter_mode = "executable"
            elif show_only_skipped.get():
                self.preview_filter_mode = "skipped"
            else:
                self.preview_filter_mode = "all"
            render_preview_rows()

        def reset_filters():
            show_only_executable.set(False)
            show_only_skipped.set(False)
            self.preview_filter_mode = "all"
            render_preview_rows()

        def render_preview_rows():
            for row_id in preview_tree.get_children():
                preview_tree.delete(row_id)

            if show_only_skipped.get():
                visible_items = [item for item in sorted_preview_items if item.get('status', '可执行') != "可执行"]
            elif show_only_executable.get():
                visible_items = [item for item in sorted_preview_items if item.get('status', '可执行') == "可执行"]
            else:
                visible_items = sorted_preview_items
            visible_items_state["items"] = visible_items

            visible_executable = sum(1 for item in visible_items if item.get('status', '可执行') == "可执行")
            visible_skipped = len(visible_items) - visible_executable
            title_var.set(f"{title_base_text}    筛选：{get_filter_scope_text()}")
            preview_summary_var.set(f"当前显示：{len(visible_items)} 项（可执行 {visible_executable}，需跳过 {visible_skipped}）")
            if button_refs["copy"] is not None:
                button_refs["copy"].configure(state=("normal" if visible_items else "disabled"))
            if button_refs["export"] is not None:
                button_refs["export"].configure(state=("normal" if visible_items else "disabled"))
            if button_refs["reason"] is not None:
                button_refs["reason"].configure(state=("normal" if visible_skipped else "disabled"))

            for item in visible_items:
                status_text = item.get('status', '可执行')
                row_id = preview_tree.insert(
                    "",
                    "end",
                    values=(item['type'], item['old_name'], item['new_name'], status_text, item.get('relative_path', item.get('path', '')))
                )
                if status_text != "可执行":
                    preview_tree.item(row_id, tags=("skipped",))

        def get_filter_scope_text():
            if show_only_skipped.get():
                return "仅看需跳过项"
            if show_only_executable.get():
                return "仅看可执行项"
            return "全部项"

        def copy_visible_items():
            visible_items = visible_items_state["items"]
            if not visible_items:
                self._show_warning_message("提示", f"当前筛选（{get_filter_scope_text()}）下没有可复制的预览项。")
                return

            text = build_visible_items_text(visible_items)
            preview_dialog.clipboard_clear()
            preview_dialog.clipboard_append(text)
            preview_dialog.update()
            self._show_info_message("完成", f"已复制当前筛选（{get_filter_scope_text()}）下的 {len(visible_items)} 条预览项到剪贴板。")

        def copy_skip_reasons():
            visible_items = visible_items_state["items"]
            skipped_items = [item for item in visible_items if item.get('status', '可执行') != "可执行"]
            if not skipped_items:
                self._show_warning_message("提示", f"当前筛选（{get_filter_scope_text()}）下没有需跳过的预览项。")
                return

            reason_counter = {}
            for item in skipped_items:
                reason = item.get('status', '未知原因')
                reason_counter[reason] = reason_counter.get(reason, 0) + 1

            lines = ["失败原因汇总\t数量"]
            for reason, count in sorted(reason_counter.items(), key=lambda pair: pair[1], reverse=True):
                lines.append(f"{reason}\t{count}")

            text = "\n".join(lines)
            preview_dialog.clipboard_clear()
            preview_dialog.clipboard_append(text)
            preview_dialog.update()
            self._show_info_message("完成", f"已复制当前筛选（{get_filter_scope_text()}）下的 {len(reason_counter)} 条失败原因汇总到剪贴板。")

        def build_visible_items_text(visible_items):
            lines = ["类型\t原名称\t新名称\t执行结果\t路径"]
            for item in visible_items:
                lines.append(
                    f"{item.get('type', '')}\t{item.get('old_name', '')}\t{item.get('new_name', '')}\t{item.get('status', '可执行')}\t{item.get('relative_path', item.get('path', ''))}"
                )
            return "\n".join(lines)

        def export_visible_items():
            visible_items = visible_items_state["items"]
            if not visible_items:
                self._show_warning_message("提示", f"当前筛选（{get_filter_scope_text()}）下没有可导出的预览项。")
                return

            default_name = f"重命名预览_{self.get_current_time().replace(':', '').replace(' ', '_')}.txt"
            save_path = filedialog.asksaveasfilename(
                title="导出预览列表",
                defaultextension=".txt",
                initialfile=default_name,
                filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")]
            )
            if not save_path:
                return

            try:
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(build_visible_items_text(visible_items))
                self._show_info_message("完成", f"已导出当前筛选（{get_filter_scope_text()}）下的 {len(visible_items)} 条预览项到:\n{save_path}")
            except Exception as e:
                self._show_error_message("错误", f"导出预览列表失败:\n{str(e)}")

        # 按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        
        button_refs["copy"] = ttk.Button(button_frame, text="复制当前列表", command=copy_visible_items)
        button_refs["copy"].pack(side=tk.LEFT)
        button_refs["reason"] = ttk.Button(button_frame, text="复制失败原因", command=copy_skip_reasons)
        button_refs["reason"].pack(side=tk.LEFT, padx=(8, 0))
        button_refs["export"] = ttk.Button(button_frame, text="导出当前列表", command=export_visible_items)
        button_refs["export"].pack(side=tk.LEFT, padx=(8, 0))
        # 20260906 205815 新增：在"关闭"按钮左侧增加"同意预览结果执行修改"按钮，点击后关闭预览对话框并触发执行回调
        if on_confirm is not None:
            def _confirm_and_close():
                preview_dialog.destroy()
                on_confirm()
            ttk.Button(button_frame, text="同意预览结果执行修改", command=_confirm_and_close).pack(side=tk.RIGHT, padx=(0, 12))
        ttk.Button(button_frame, text="关闭", command=preview_dialog.destroy).pack(side=tk.RIGHT)
        render_preview_rows()

        # 20260906 205815 新增：动态计算预览对话框尺寸，确保打开后能直接看到最下方按钮排（复制/导出/关闭）
        preview_dialog.update_idletasks()
        req_width = main_frame.winfo_reqwidth() + 30
        req_height = main_frame.winfo_reqheight() + 30
        dialog_width = max(700, min(req_width, preview_dialog.winfo_screenwidth() - 100))
        dialog_height = max(450, min(req_height, int(preview_dialog.winfo_screenheight() * 0.8)))
        x = (preview_dialog.winfo_screenwidth() - dialog_width) // 2
        y = max(50, (preview_dialog.winfo_screenheight() - dialog_height) // 3)
        preview_dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")

def main():
    root = tk.Tk()
    app = DirCopyApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
